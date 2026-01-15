"""
Create Session Summary Document in Word format
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_session_summary():
    doc = Document()

    # Title
    title = doc.add_heading('Session Summary - November 22, 2025', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Overview
    doc.add_heading('Overview', level=1)
    doc.add_paragraph(
        'This session focused on creating the weekly data infrastructure for ALL US stocks (~20,000) '
        'and cryptos, scheduler monitoring for the Admin panel, and NLP search capabilities. '
        'All changes are incremental and designed to work alongside existing GCP Cloud Run services.'
    )

    # BigQuery Tables
    doc.add_heading('1. BigQuery Tables Created', level=1)

    doc.add_heading('stocks_weekly_summary (58 columns)', level=2)
    doc.add_paragraph('Weekly summary for ALL US stocks including:')
    bullets = [
        'Basic info: symbol, name, exchange, type, sector, industry',
        'Price data: current, open, high, low, close, previous_close',
        'Performance: weekly/monthly/YTD change percentages',
        'Valuation: market_cap, PE ratio, EPS, dividend yield, beta',
        'Trading ranges: 52-week high/low, percent from high/low',
        'Volatility: weekly/monthly volatility, ATR, ATR percent',
        'Day trading scores: day_trade_score, liquidity_score, momentum_score',
        'Technical indicators: RSI, MACD, SMAs',
        'Classifications: market_cap_category, volatility_category, momentum_category',
        'Active list flags: is_active_pick, active_pick_reason'
    ]
    for item in bullets:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('cryptos_weekly_summary (55 columns)', level=2)
    doc.add_paragraph('Weekly summary for ALL cryptos with same structure plus crypto-specific categories: DeFi, Meme, Smart Contract Platform, Layer 2, etc.')

    doc.add_heading('active_trading_list (16 columns)', level=2)
    doc.add_paragraph('Top picks for day trading - generated from weekly analysis containing top 100 high-volatility stocks and cryptos.')

    doc.add_heading('scheduler_execution_log (30 columns)', level=2)
    doc.add_paragraph('Tracks all scheduler runs including execution timing, status, statistics, and API usage.')

    # Cloud Functions
    doc.add_heading('2. Cloud Functions Created', level=1)

    doc.add_heading('cloud_function_weekly_stocks/', level=2)
    doc.add_paragraph('Fetches ALL US stocks from Twelve Data API with features:')
    funcs = [
        'Gets complete US stock list from Twelve Data',
        'Fetches quote, weekly data, and statistics for each stock',
        'Calculates day trading scores, volatility metrics',
        'Logs execution to scheduler_execution_log',
        'Batched uploads to BigQuery'
    ]
    for item in funcs:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('cloud_function_weekly_cryptos/', level=2)
    doc.add_paragraph('Fetches ALL cryptos from Twelve Data API with same processing and category mapping.')

    # Scheduler Configuration
    doc.add_heading('3. Scheduler Configuration', level=1)

    table = doc.add_table(rows=9, cols=4)
    table.style = 'Table Grid'

    headers = ['Scheduler', 'Schedule', 'Table', 'Description']
    header_row = table.rows[0].cells
    for i, header in enumerate(headers):
        header_row[i].text = header

    schedulers = [
        ('daily-crypto-fetcher', '12:00 AM ET', 'crypto_daily', 'Daily crypto OHLC'),
        ('hourly-crypto-fetcher', 'Every hour', 'crypto_hourly_data', 'Hourly crypto'),
        ('fivemin-top10-fetcher', 'Every 5 min', 'crypto_5min_top10_gainers', 'Top 10 gainers'),
        ('daily-stock-fetcher', '6:00 PM ET', 'stocks_daily', 'Daily stock OHLC'),
        ('hourly-stock-fetcher', 'Every hour', 'stocks_hourly_data', 'Hourly stocks'),
        ('fivemin-stock-fetcher', 'Every 5 min', 'stocks_5min_data', '5-min stocks'),
        ('weekly-stock-fetcher', 'Sunday 11 PM ET', 'stocks_weekly_summary', 'ALL US stocks (NEW)'),
        ('weekly-crypto-fetcher', 'Sunday 10 PM ET', 'cryptos_weekly_summary', 'ALL cryptos (NEW)')
    ]

    for i, sched in enumerate(schedulers):
        row = table.rows[i + 1].cells
        for j, val in enumerate(sched):
            row[j].text = val

    # API Endpoints
    doc.add_heading('4. API Endpoints Added (Incremental)', level=1)
    doc.add_paragraph('These endpoints are ADDED to the existing crypto-trading-api without modifying existing endpoints:')

    doc.add_heading('Admin Endpoints:', level=2)
    endpoints = [
        'GET /api/admin/scheduler-status - Get scheduler status and logs',
        'POST /api/admin/trigger-scheduler - Manually trigger a scheduler'
    ]
    for ep in endpoints:
        doc.add_paragraph(ep, style='List Bullet')

    doc.add_heading('Analysis Endpoints:', level=2)
    analysis_eps = [
        'GET /api/analysis/stocks/weekly - Weekly stock analysis with filters',
        'GET /api/analysis/cryptos/weekly - Weekly crypto analysis with filters',
        'GET /api/analysis/active-list - Get top picks for day trading',
        'POST /api/analysis/nlp-search - Natural language search'
    ]
    for ep in analysis_eps:
        doc.add_paragraph(ep, style='List Bullet')

    # NLP Search
    doc.add_heading('5. NLP Search Capability', level=1)
    doc.add_paragraph('Supports natural language queries like:')
    queries = [
        '"high volatility tech stocks"',
        '"bullish meme cryptos"',
        '"large cap financial stocks for day trading"',
        '"stable healthcare stocks"'
    ]
    for q in queries:
        doc.add_paragraph(q, style='List Bullet')

    # Frontend Components
    doc.add_heading('6. Frontend Components', level=1)
    doc.add_heading('SchedulerMonitoring.jsx', level=2)
    doc.add_paragraph('New Admin component for monitoring all 8 data pipelines:')
    features = [
        'Shows last execution time and status for each scheduler',
        'Duration, success rate, total runs (30 days)',
        'Recent execution logs table',
        'Manual trigger capability for each scheduler'
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')

    # Files Created
    doc.add_heading('7. Files Created/Modified', level=1)

    doc.add_heading('New Files:', level=2)
    new_files = [
        'create_weekly_stocks_table.py - Creates weekly tables in BigQuery',
        'create_scheduler_log_table.py - Creates scheduler monitoring tables',
        'cloud_function_weekly_stocks/main.py - Weekly stock loader',
        'cloud_function_weekly_stocks/requirements.txt',
        'cloud_function_weekly_cryptos/main.py - Weekly crypto loader',
        'cloud_function_weekly_cryptos/requirements.txt',
        'deploy_weekly_functions.py - Deployment script',
        'stock-price-app/src/components/SchedulerMonitoring.jsx - Admin UI'
    ]
    for f in new_files:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_heading('Modified Files (Incremental changes only):', level=2)
    mod_files = [
        'cloud_function_api/main.py - Added new endpoints (existing endpoints unchanged)',
        'stock-price-app/src/services/api.js - Added new API methods'
    ]
    for f in mod_files:
        doc.add_paragraph(f, style='List Bullet')

    # Deployment Steps
    doc.add_heading('8. Deployment Steps (Incremental)', level=1)
    doc.add_paragraph('All deployments are incremental and will not affect existing running services:')

    steps = [
        'Step 1: Deploy updated API (adds new endpoints, keeps existing):',
        '  gcloud run deploy crypto-trading-api --source cloud_function_api --region us-central1 --project cryptobot-462709',
        '',
        'Step 2: Deploy weekly stock function (new service):',
        '  gcloud functions deploy weekly-stock-fetcher --gen2 --runtime=python312 --region=us-central1 --source=cloud_function_weekly_stocks ...',
        '',
        'Step 3: Deploy weekly crypto function (new service):',
        '  gcloud functions deploy weekly-crypto-fetcher --gen2 --runtime=python312 --region=us-central1 --source=cloud_function_weekly_cryptos ...',
        '',
        'Step 4: Create weekly schedulers (new jobs):',
        '  gcloud scheduler jobs create http weekly-stock-fetch-job --schedule="0 23 * * 0" ...',
        '  gcloud scheduler jobs create http weekly-crypto-fetch-job --schedule="0 22 * * 0" ...',
        '',
        'Step 5: Rebuild and deploy frontend (includes new SchedulerMonitoring component):',
        '  gcloud run deploy crypto-trading-app --source stock-price-app --region us-central1 --project cryptobot-462709'
    ]
    for step in steps:
        doc.add_paragraph(step)

    # Document Reference
    doc.add_heading('9. Document Reference', level=1)
    doc.add_paragraph('TOP_100_STOCKS_CRYPTOS.docx - Contains top 100 stocks and cryptos list with prices')

    # Save
    doc.save('SESSION_SUMMARY_NOV22.docx')
    print("Created: SESSION_SUMMARY_NOV22.docx")

if __name__ == "__main__":
    create_session_summary()
