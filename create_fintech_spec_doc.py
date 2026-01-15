"""
Create professional Fintech AI Data Architecture Specification Document
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import sys
import io

if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

doc = Document()

# Set document styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_heading('AIAlgoTradeHits.com', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Fintech AI Data Architecture & Implementation Specification')
run.bold = True
run.font.size = Pt(16)

# Metadata
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Master Technical Document').bold = True
meta.add_run('\n')
meta.add_run('Version 2.0 | November 2025')
meta.add_run('\n')
meta.add_run('Platform: Google Cloud Platform (GCP)')
meta.add_run('\n')
meta.add_run('AI Engine: Vertex AI + Gemini 3 Pro')

doc.add_page_break()

# Table of Contents
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. Multi-Layer Data Architecture',
    '3. BigQuery Table Structure & Schema',
    '4. Phase 1 Missing Fields Analysis',
    '5. Vertex AI & Gemini 3 Integration',
    '6. RAG & Document Intelligence',
    '7. Implementation Phases',
    '8. GCP Infrastructure & Costs',
    '9. Technical Indicator Formulas',
    '10. Implementation Roadmap'
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# Executive Summary
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'AIAlgoTradeHits.com is a comprehensive cross-asset, AI-driven analytics, alerting, and intelligence platform '
    'delivering institutional-grade insights across equities, ETFs, cryptocurrency, forex, commodities, and '
    'interest-rate markets. Built entirely on Google Cloud Platform, the system leverages Vertex AI, Gemini 3 LLM, '
    'BigQuery, and Cloud Run to create a next-generation fintech trading intelligence engine.'
)

doc.add_heading('Platform Vision', level=2)
doc.add_paragraph(
    'The system integrates several advanced components that together form a complete AI market analysis engine:'
)
vision_points = [
    'Clean, normalized market data derived from multi-asset providers through a high-integrity, multi-layer data architecture',
    'Deterministic indicators and engineered features that transform raw OHLCV into structured intelligence',
    'Machine learning models capable of understanding regimes, volatility, reversals, and trend transitions',
    'LLM-powered explanations (Gemini 3 + RAG) generating research-style commentary grounded in actual data',
    'Real-time serving infrastructure for alerts, dashboards, and chart overlays',
    'Compliance and audit guardrails suitable for banks, regulated entities, and enterprise use'
]
for point in vision_points:
    doc.add_paragraph(point, style='List Bullet')

# Platform Components Table
doc.add_heading('Technology Stack', level=2)
table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Component'
hdr[1].text = 'Technology'
hdr[2].text = 'Purpose'
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True

data = [
    ('Data Warehouse', 'BigQuery', 'Multi-layer data architecture (Bronze/Silver/Gold)'),
    ('AI/ML Engine', 'Vertex AI', 'Predictive models, pattern recognition, regime classification'),
    ('LLM Intelligence', 'Gemini 3 Pro', 'Natural language analysis, RAG-powered commentary'),
    ('Document Intelligence', 'Document AI + RAG', 'Research document processing, knowledge retrieval'),
    ('Real-Time Serving', 'Cloud Run + Cloud Functions', 'API endpoints, scheduled data collection'),
    ('Frontend', 'React + TradingView Charts', 'Professional trading interface'),
]
for i, row_data in enumerate(data):
    row = table.rows[i + 1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

# Part 2: Data Architecture
doc.add_page_break()
doc.add_heading('2. Multi-Layer Data Architecture', level=1)

doc.add_paragraph(
    'The platform implements a rigorous four-layer data architecture inspired by institutional quant trading desks. '
    'This architecture ensures traceability, reproducibility, auditability, real-time serving, LLM grounding, '
    'model consistency, and institutional compliance.'
)

doc.add_heading('2.1 Bronze Layer - Raw Market Data', level=2)
doc.add_paragraph(
    'Stores raw, immutable payloads exactly as received from providers (TwelveData, Kraken, Finnhub) without '
    'corrections or transformations. This layer provides:'
)
bronze_points = [
    'Exact API responses preserved',
    'Full lineage and audit trail',
    'Compliance-grade immutability',
    'Replay capability for debugging',
    'Multi-provider consolidation'
]
for point in bronze_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_heading('2.2 Silver Layer - Cleaned & Standardized Data', level=2)
doc.add_paragraph(
    'Data becomes trusted, normalized, and analytic-ready. This is where data quality is enforced:'
)
silver_points = [
    'Normalized timestamps (UTC)',
    'Missing data handling and repair',
    'Anomaly removal (zero-volume, spikes)',
    'Session calendar alignment',
    'Provider discrepancy resolution',
    'Corporate actions application (stocks)'
]
for point in silver_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_heading('2.3 Gold Layer - Engineered Features & Intelligence', level=2)
doc.add_paragraph(
    'Converts cleaned OHLCV into ML-ready intelligence. This is the analytical backbone:'
)
gold_points = [
    '70+ Technical Indicators (momentum, trend, volatility, volume)',
    'Pattern Recognition Signals (head & shoulders, triangles, flags)',
    'Cross-Asset Correlations (BTC-ETH, SPY-QQQ, USD strength)',
    'Volatility Regime Classification',
    'ML-Ready Feature Vectors',
    'Sentiment integration'
]
for point in gold_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_heading('2.4 Wide Tables - Fast-Access Serving Layer', level=2)
doc.add_paragraph(
    'Materialized flattened tables optimized for real-time access:'
)
wide_points = [
    'One row per symbol per timestamp with all features',
    'Real-time UI/API access optimized',
    'LLM RAG retrieval optimized',
    'ML inference ready',
    'Dashboard and alert engine feeds'
]
for point in wide_points:
    doc.add_paragraph(point, style='List Bullet')

# BigQuery Tables
doc.add_page_break()
doc.add_heading('3. BigQuery Table Structure & Schema', level=1)

doc.add_heading('3.1 Asset Class Tables', level=2)
table2 = doc.add_table(rows=9, cols=4)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Table Name'
hdr2[1].text = 'Asset Class'
hdr2[2].text = 'Timeframes'
hdr2[3].text = 'Data Source'
for cell in hdr2:
    cell.paragraphs[0].runs[0].bold = True

tables_data = [
    ('stocks_historical_daily', 'US Equities', 'Daily, Weekly', 'TwelveData'),
    ('stocks_hourly', 'US Equities', 'Hourly, 5-min', 'TwelveData'),
    ('cryptos_historical_daily', 'Cryptocurrency', 'Daily, Weekly', 'Kraken/TwelveData'),
    ('crypto_hourly_data', 'Cryptocurrency', 'Hourly, 5-min', 'Kraken'),
    ('etfs_historical_daily', 'ETFs', 'Daily, Weekly', 'TwelveData'),
    ('forex_historical_daily', 'FX Pairs', 'Daily, Weekly', 'TwelveData'),
    ('indices_historical_daily', 'Market Indices', 'Daily, Weekly', 'TwelveData'),
    ('commodities_historical_daily', 'Commodities', 'Daily, Weekly', 'TwelveData'),
]
for i, row_data in enumerate(tables_data):
    row = table2.rows[i + 1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]
    row[3].text = row_data[3]

# Schema sections
doc.add_heading('3.2 Core OHLCV Fields (Required)', level=2)
ohlcv_table = doc.add_table(rows=10, cols=4)
ohlcv_table.style = 'Table Grid'
ohlcv_hdr = ohlcv_table.rows[0].cells
ohlcv_hdr[0].text = 'Field Name'
ohlcv_hdr[1].text = 'Type'
ohlcv_hdr[2].text = 'Description'
ohlcv_hdr[3].text = 'Status'
for cell in ohlcv_hdr:
    cell.paragraphs[0].runs[0].bold = True

ohlcv_fields = [
    ('symbol', 'STRING', 'Asset ticker symbol', 'Available'),
    ('datetime', 'TIMESTAMP', 'Candle timestamp (UTC)', 'Available'),
    ('open', 'FLOAT64', 'Opening price', 'Available'),
    ('high', 'FLOAT64', 'High price', 'Available'),
    ('low', 'FLOAT64', 'Low price', 'Available'),
    ('close', 'FLOAT64', 'Closing price', 'Available'),
    ('volume', 'FLOAT64', 'Trading volume', 'Available'),
    ('candle_body_pct', 'FLOAT64', '(close - open) / open', 'ADD'),
    ('candle_range_pct', 'FLOAT64', '(high - low) / open', 'ADD'),
]
for i, row_data in enumerate(ohlcv_fields):
    row = ohlcv_table.rows[i + 1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]
    row[3].text = row_data[3]

# Missing Fields Section
doc.add_page_break()
doc.add_heading('4. Phase 1 Missing Fields Analysis', level=1)
doc.add_paragraph(
    'Based on the Phase 1 Methodology specification, the following fields need to be added to BigQuery tables '
    'for optimal ML training. Fields are prioritized based on ML impact and implementation complexity.'
)

doc.add_heading('4.1 Priority 1: Critical ML Features (Add Immediately)', level=2)
doc.add_paragraph('These fields are essential for core ML model performance:')

p1_table = doc.add_table(rows=16, cols=4)
p1_table.style = 'Table Grid'
p1_hdr = p1_table.rows[0].cells
p1_hdr[0].text = 'Field Name'
p1_hdr[1].text = 'Type'
p1_hdr[2].text = 'Description'
p1_hdr[3].text = 'Formula'
for cell in p1_hdr:
    cell.paragraphs[0].runs[0].bold = True

p1_fields = [
    ('weekly_log_return', 'FLOAT64', 'Log return', 'ln(close_t/close_{t-1})'),
    ('return_2w', 'FLOAT64', '2-week return', 'close_t/close_{t-2} - 1'),
    ('return_4w', 'FLOAT64', '4-week return', 'close_t/close_{t-4} - 1'),
    ('rsi_slope', 'FLOAT64', 'RSI change rate', 'rsi_t - rsi_{t-1}'),
    ('rsi_zscore', 'FLOAT64', 'RSI z-score', '(rsi - mean_100w) / std_100w'),
    ('rsi_overbought_flag', 'INT64', 'RSI > 70 flag', '1 if rsi > 70 else 0'),
    ('rsi_oversold_flag', 'INT64', 'RSI < 30 flag', '1 if rsi < 30 else 0'),
    ('macd_hist', 'FLOAT64', 'MACD histogram', 'macd - macd_signal'),
    ('macd_cross_flag', 'INT64', 'Cross signal', '+1 bull, -1 bear, 0 none'),
    ('ema_20', 'FLOAT64', 'EMA 20-period', 'EMA(close, 20)'),
    ('ema_50', 'FLOAT64', 'EMA 50-period', 'EMA(close, 50)'),
    ('ema_200', 'FLOAT64', 'EMA 200-period', 'EMA(close, 200)'),
    ('close_vs_sma20_pct', 'FLOAT64', 'Price vs SMA20', '(close/SMA20 - 1) * 100'),
    ('ema_slope_20', 'FLOAT64', 'EMA20 slope', 'ema20_t - ema20_{t-1}'),
    ('ema_slope_50', 'FLOAT64', 'EMA50 slope', 'ema50_t - ema50_{t-1}'),
]
for i, row_data in enumerate(p1_fields):
    row = p1_table.rows[i + 1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]
    row[3].text = row_data[3]

doc.add_heading('4.2 Priority 2: Enhanced Features (Next Sprint)', level=2)
p2_table = doc.add_table(rows=10, cols=4)
p2_table.style = 'Table Grid'
p2_hdr = p2_table.rows[0].cells
p2_hdr[0].text = 'Field Name'
p2_hdr[1].text = 'Type'
p2_hdr[2].text = 'Description'
p2_hdr[3].text = 'Formula'
for cell in p2_hdr:
    cell.paragraphs[0].runs[0].bold = True

p2_fields = [
    ('atr_pct', 'FLOAT64', 'ATR as % of price', '(ATR14 / close) * 100'),
    ('atr_zscore', 'FLOAT64', 'ATR z-score', '(atr - mean) / std'),
    ('atr_slope', 'FLOAT64', 'ATR change rate', 'atr_t - atr_{t-1}'),
    ('volume_zscore', 'FLOAT64', 'Volume z-score', '(vol - mean_20) / std_20'),
    ('volume_ratio', 'FLOAT64', 'Volume vs MA', 'volume / SMA(volume, 20)'),
    ('adx', 'FLOAT64', 'ADX (14)', 'Standard ADX calculation'),
    ('plus_di', 'FLOAT64', '+DI (14)', 'Directional indicator +'),
    ('minus_di', 'FLOAT64', '-DI (14)', 'Directional indicator -'),
    ('candle_body_pct', 'FLOAT64', 'Body % of range', '(close-open) / open'),
]
for i, row_data in enumerate(p2_fields):
    row = p2_table.rows[i + 1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]
    row[3].text = row_data[3]

doc.add_heading('4.3 Priority 3: Advanced Structure (Phase 2)', level=2)
p3_table = doc.add_table(rows=7, cols=4)
p3_table.style = 'Table Grid'
p3_hdr = p3_table.rows[0].cells
p3_hdr[0].text = 'Field Name'
p3_hdr[1].text = 'Type'
p3_hdr[2].text = 'Description'
p3_hdr[3].text = 'Purpose'
for cell in p3_hdr:
    cell.paragraphs[0].runs[0].bold = True

p3_fields = [
    ('pivot_high_flag', 'INT64', 'Pivot high detected', 'Structure detection'),
    ('pivot_low_flag', 'INT64', 'Pivot low detected', 'Structure detection'),
    ('dist_to_pivot_high', 'FLOAT64', 'Distance to last pivot high', 'Support/Resistance'),
    ('dist_to_pivot_low', 'FLOAT64', 'Distance to last pivot low', 'Support/Resistance'),
    ('regime_state', 'INT64', 'Numeric regime (1-6)', 'ML classification'),
    ('regime_confidence', 'FLOAT64', 'Regime confidence (0-1)', 'Model certainty'),
]
for i, row_data in enumerate(p3_fields):
    row = p3_table.rows[i + 1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]
    row[3].text = row_data[3]

# Part 5: AI Architecture
doc.add_page_break()
doc.add_heading('5. Vertex AI & Gemini 3 Integration', level=1)

doc.add_heading('5.1 Vertex AI Model Types', level=2)
ai_table = doc.add_table(rows=6, cols=3)
ai_table.style = 'Table Grid'
ai_hdr = ai_table.rows[0].cells
ai_hdr[0].text = 'Model Type'
ai_hdr[1].text = 'Vertex AI Service'
ai_hdr[2].text = 'Use Case'
for cell in ai_hdr:
    cell.paragraphs[0].runs[0].bold = True

ai_data = [
    ('Gradient Boosting', 'AutoML Tables', 'Regime classification, direction prediction'),
    ('Time Series', 'Vertex AI Forecasting', 'Price prediction, volatility forecasting'),
    ('LSTM/Transformer', 'Custom Training', 'Sequential pattern learning'),
    ('Vision Models', 'AutoML Vision', 'Chart pattern recognition'),
    ('LLM', 'Gemini 3 Pro', 'Commentary generation, analysis'),
]
for i, row_data in enumerate(ai_data):
    row = ai_table.rows[i + 1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

doc.add_heading('5.2 Gemini 3 LLM Use Cases', level=2)
llm_table = doc.add_table(rows=6, cols=4)
llm_table.style = 'Table Grid'
llm_hdr = llm_table.rows[0].cells
llm_hdr[0].text = 'Use Case'
llm_hdr[1].text = 'Input'
llm_hdr[2].text = 'Output'
llm_hdr[3].text = 'Frequency'
for cell in llm_hdr:
    cell.paragraphs[0].runs[0].bold = True

llm_data = [
    ('Market Commentary', 'Indicators, patterns, regime', 'Research-style analysis', 'Real-time'),
    ('Alert Explanations', 'Alert trigger data', 'Natural language reasoning', 'Per alert'),
    ('Pattern Analysis', 'Chart patterns', 'Professional interpretation', 'On detection'),
    ('Risk Assessment', 'Cross-asset correlations', 'Risk narrative', 'Hourly'),
    ('Q&A Interface', 'User questions + RAG', 'Informed responses', 'On-demand'),
]
for i, row_data in enumerate(llm_data):
    row = llm_table.rows[i + 1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]
    row[3].text = row_data[3]

# RAG Section
doc.add_page_break()
doc.add_heading('6. RAG & Document Intelligence', level=1)

doc.add_heading('6.1 RAG Architecture Overview', level=2)
doc.add_paragraph(
    'The RAG (Retrieval-Augmented Generation) pipeline combines multiple data sources to provide grounded, '
    'accurate responses from Gemini 3. This prevents hallucination and ensures all LLM outputs are backed by '
    'actual data sources.'
)

doc.add_heading('RAG Components:', level=3)
rag_components = [
    'BigQuery Market Data: Real-time OHLCV, indicators, and ML predictions',
    'Vector Database: Research document embeddings using Vertex AI Embeddings',
    'Document Store: PDF and DOCX research documents via Document AI',
    'Knowledge Graph: Entity relationships for cross-asset understanding'
]
for comp in rag_components:
    doc.add_paragraph(comp, style='List Bullet')

doc.add_heading('6.2 Document AI Integration', level=2)
docai_table = doc.add_table(rows=5, cols=3)
docai_table.style = 'Table Grid'
docai_hdr = docai_table.rows[0].cells
docai_hdr[0].text = 'Component'
docai_hdr[1].text = 'GCP Service'
docai_hdr[2].text = 'Purpose'
for cell in docai_hdr:
    cell.paragraphs[0].runs[0].bold = True

docai_data = [
    ('PDF Processor', 'Document AI', 'Extract text from research PDFs'),
    ('Embedding Generator', 'Vertex AI Embeddings', 'Create vector representations'),
    ('Vector Store', 'Vertex AI Vector Search', 'Semantic search over documents'),
    ('Knowledge Graph', 'Enterprise Knowledge Graph', 'Entity relationships'),
]
for i, row_data in enumerate(docai_data):
    row = docai_table.rows[i + 1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

# Implementation Phases
doc.add_page_break()
doc.add_heading('7. Implementation Phases', level=1)

doc.add_heading('7.1 Phase 1: Trader-Focused AI Alert Engine', level=2)
doc.add_paragraph('Timeline: Current - 3 months')
doc.add_paragraph('Focus: Real-time alerts for equities and cryptocurrency')

doc.add_heading('MVP Symbol Universe:', level=3)
symbols_table = doc.add_table(rows=7, cols=3)
symbols_table.style = 'Table Grid'
sym_hdr = symbols_table.rows[0].cells
sym_hdr[0].text = 'Category'
sym_hdr[1].text = 'Examples'
sym_hdr[2].text = 'Count'
for cell in sym_hdr:
    cell.paragraphs[0].runs[0].bold = True

sym_data = [
    ('US Equities', 'AAPL, NVDA, JPM, MSFT, GOOGL, AMZN, META, TSLA', '50'),
    ('ETFs', 'SPY, QQQ, IWM, DIA, VTI, VOO, GLD, SLV', '30'),
    ('Cryptocurrency', 'BTC/USD, ETH/USD, SOL/USD, XRP/USD, ADA/USD', '20'),
    ('Forex', 'EUR/USD, GBP/USD, USD/JPY, EUR/GBP, AUD/USD', '20'),
    ('Indices', 'SPX, NDX, DJI, VIX, DAX, FTSE', '14'),
    ('Commodities', 'XAU/USD, XAG/USD, CL, NG, ZC', '16'),
]
for i, row_data in enumerate(sym_data):
    row = symbols_table.rows[i + 1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

doc.add_heading('Phase 1 Deliverables:', level=3)
p1_deliverables = [
    'Real-Time Alert System (trend flips, momentum reversals, breakouts)',
    'AI-Enhanced Dashboards (regime classification, volatility heatmaps)',
    'LLM Commentary (alert explanations, daily summaries)',
    'Pattern Recognition (head & shoulders, triangles, flags)',
    'Multi-timeframe analysis (daily, hourly, 5-minute)'
]
for d in p1_deliverables:
    doc.add_paragraph(d, style='List Bullet')

doc.add_heading('7.2 Phase 2: Institutional Treasury Intelligence', level=2)
doc.add_paragraph('Timeline: 3-6 months')
doc.add_paragraph('Focus: Enterprise-grade macro and treasury intelligence')

doc.add_heading('Phase 2 Capabilities:', level=3)
p2_caps = [
    'Cross-Asset Correlation Engine',
    'Macro Regime Classification (risk-on/off detection)',
    'Treasury Exposure Analytics',
    'Yield Curve Analytics',
    'Early Warning Risk Systems',
    'AI-Generated Institutional Reports',
    'Compliance and Audit Logging'
]
for cap in p2_caps:
    doc.add_paragraph(cap, style='List Bullet')

# Cost Analysis
doc.add_page_break()
doc.add_heading('8. GCP Infrastructure & Costs', level=1)

doc.add_heading('8.1 Service Architecture', level=2)
doc.add_paragraph('GCP Project: cryptobot-462709')

services_table = doc.add_table(rows=7, cols=3)
services_table.style = 'Table Grid'
svc_hdr = services_table.rows[0].cells
svc_hdr[0].text = 'Service'
svc_hdr[1].text = 'Purpose'
svc_hdr[2].text = 'Configuration'
for cell in svc_hdr:
    cell.paragraphs[0].runs[0].bold = True

svc_data = [
    ('Cloud Run', 'Frontend + API hosting', 'Auto-scaling, min 0 instances'),
    ('Cloud Functions', 'Data collection (6 functions)', 'Scheduled triggers'),
    ('BigQuery', 'Data warehouse', '6 asset class tables'),
    ('Cloud Scheduler', 'Job triggers', 'Daily, hourly, 5-min'),
    ('Vertex AI', 'ML models', 'AutoML + Custom training'),
    ('Cloud Storage', 'Document storage', 'Research PDFs, exports'),
]
for i, row_data in enumerate(svc_data):
    row = services_table.rows[i + 1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

doc.add_heading('8.2 Cost Optimization', level=2)
cost_table = doc.add_table(rows=7, cols=3)
cost_table.style = 'Table Grid'
cost_hdr = cost_table.rows[0].cells
cost_hdr[0].text = 'Service'
cost_hdr[1].text = 'Optimization Strategy'
cost_hdr[2].text = 'Monthly Est.'
for cell in cost_hdr:
    cell.paragraphs[0].runs[0].bold = True

cost_data = [
    ('BigQuery', 'Partitioned tables, clustering by symbol', '$50-100'),
    ('Cloud Functions', 'Efficient scheduling, cold start optimization', '$130'),
    ('Cloud Run', 'Auto-scaling, minimum instances = 0', '$50-100'),
    ('Vertex AI', 'Batch predictions, model caching', '$200-500'),
    ('Gemini 3', 'Response caching, prompt optimization', '$100-300'),
    ('TOTAL', '', '$530-1,130/month'),
]
for i, row_data in enumerate(cost_data):
    row = cost_table.rows[i + 1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

# Roadmap
doc.add_page_break()
doc.add_heading('10. Implementation Roadmap', level=1)

roadmap_table = doc.add_table(rows=9, cols=3)
roadmap_table.style = 'Table Grid'
road_hdr = roadmap_table.rows[0].cells
road_hdr[0].text = 'Week'
road_hdr[1].text = 'Task'
road_hdr[2].text = 'Deliverables'
for cell in road_hdr:
    cell.paragraphs[0].runs[0].bold = True

road_data = [
    ('1-2', 'Data Schema Enhancement', 'Add Priority 1 fields to BigQuery'),
    ('2-3', 'Data Backfill', 'Recalculate indicators for historical data'),
    ('3-4', 'Vertex AI Setup', 'Configure Feature Store, training pipelines'),
    ('4-5', 'ML Model Training', 'Train baseline models (regime, direction)'),
    ('5-6', 'LLM Integration', 'RAG pipeline, Gemini 3 prompts'),
    ('6-7', 'Alert System', 'Real-time alert generation'),
    ('7-8', 'Testing', 'Backtest, evaluate, optimize'),
    ('8+', 'Phase 2 Planning', 'Treasury intelligence roadmap'),
]
for i, row_data in enumerate(road_data):
    row = roadmap_table.rows[i + 1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

# Save document
doc.save('FINTECH_AI_DATA_ARCHITECTURE_MASTER_SPECIFICATION.docx')
print('Document created successfully: FINTECH_AI_DATA_ARCHITECTURE_MASTER_SPECIFICATION.docx')
