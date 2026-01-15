"""
Create Top 100 Stocks and Cryptos Word Document
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_top_100_doc():
    doc = Document()

    # Title
    title = doc.add_heading('Top 100 Stocks and Cryptos for Day Trading', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
    doc.add_paragraph('High volatility assets suitable for day trading based on volume, price movement, and liquidity.')

    # ============ TOP 50 STOCKS ============
    doc.add_heading('Top 50 US Stocks for Day Trading', level=1)
    doc.add_paragraph('Selected based on: High daily volume, volatility, tight spreads, and strong momentum.')

    # Create stocks table
    stocks_table = doc.add_table(rows=51, cols=6)
    stocks_table.style = 'Table Grid'

    # Header
    headers = ['#', 'Symbol', 'Company Name', 'Sector', 'Avg Volume', 'Why Selected']
    for i, header in enumerate(headers):
        stocks_table.rows[0].cells[i].text = header

    # Top 50 day trading stocks (high volume, volatile)
    top_stocks = [
        ('1', 'TSLA', 'Tesla Inc', 'Consumer Cyclical', '100M+', 'High volatility, heavy volume'),
        ('2', 'NVDA', 'NVIDIA Corp', 'Technology', '50M+', 'AI leader, strong momentum'),
        ('3', 'AMD', 'Advanced Micro Devices', 'Technology', '60M+', 'Chip sector volatility'),
        ('4', 'AAPL', 'Apple Inc', 'Technology', '80M+', 'Most liquid stock'),
        ('5', 'AMZN', 'Amazon.com Inc', 'Consumer Cyclical', '40M+', 'E-commerce leader'),
        ('6', 'META', 'Meta Platforms', 'Technology', '25M+', 'Social media giant'),
        ('7', 'MSFT', 'Microsoft Corp', 'Technology', '30M+', 'Cloud/AI growth'),
        ('8', 'GOOGL', 'Alphabet Inc', 'Technology', '25M+', 'Search/AI dominance'),
        ('9', 'SPY', 'S&P 500 ETF', 'ETF', '100M+', 'Market benchmark'),
        ('10', 'QQQ', 'Nasdaq 100 ETF', 'ETF', '50M+', 'Tech-heavy ETF'),
        ('11', 'SOXL', 'Semiconductor Bull 3X', 'ETF', '30M+', 'Leveraged chip play'),
        ('12', 'TQQQ', 'Nasdaq Bull 3X', 'ETF', '40M+', 'Leveraged tech'),
        ('13', 'PLTR', 'Palantir Technologies', 'Technology', '50M+', 'AI/Data analytics'),
        ('14', 'MARA', 'Marathon Digital', 'Crypto/Mining', '30M+', 'Bitcoin proxy'),
        ('15', 'RIOT', 'Riot Platforms', 'Crypto/Mining', '25M+', 'Crypto miner'),
        ('16', 'COIN', 'Coinbase Global', 'Financial', '15M+', 'Crypto exchange'),
        ('17', 'SOFI', 'SoFi Technologies', 'Financial', '40M+', 'Fintech growth'),
        ('18', 'NIO', 'NIO Inc', 'Consumer Cyclical', '30M+', 'EV maker China'),
        ('19', 'RIVN', 'Rivian Automotive', 'Consumer Cyclical', '20M+', 'EV startup'),
        ('20', 'LCID', 'Lucid Group', 'Consumer Cyclical', '25M+', 'Luxury EV'),
        ('21', 'F', 'Ford Motor', 'Consumer Cyclical', '60M+', 'EV transition'),
        ('22', 'GM', 'General Motors', 'Consumer Cyclical', '15M+', 'Auto giant'),
        ('23', 'BAC', 'Bank of America', 'Financial', '40M+', 'Banking leader'),
        ('24', 'JPM', 'JPMorgan Chase', 'Financial', '10M+', 'Largest US bank'),
        ('25', 'C', 'Citigroup', 'Financial', '15M+', 'Global bank'),
        ('26', 'WFC', 'Wells Fargo', 'Financial', '20M+', 'Regional banking'),
        ('27', 'XOM', 'Exxon Mobil', 'Energy', '20M+', 'Oil major'),
        ('28', 'CVX', 'Chevron Corp', 'Energy', '10M+', 'Energy giant'),
        ('29', 'OXY', 'Occidental Petroleum', 'Energy', '15M+', 'Buffett favorite'),
        ('30', 'SLB', 'Schlumberger', 'Energy', '12M+', 'Oil services'),
        ('31', 'INTC', 'Intel Corp', 'Technology', '50M+', 'Chip turnaround'),
        ('32', 'MU', 'Micron Technology', 'Technology', '20M+', 'Memory chips'),
        ('33', 'QCOM', 'Qualcomm', 'Technology', '10M+', 'Mobile chips'),
        ('34', 'AVGO', 'Broadcom Inc', 'Technology', '5M+', 'AI infrastructure'),
        ('35', 'CRM', 'Salesforce', 'Technology', '8M+', 'Cloud software'),
        ('36', 'SNOW', 'Snowflake', 'Technology', '5M+', 'Data cloud'),
        ('37', 'SHOP', 'Shopify', 'Technology', '8M+', 'E-commerce platform'),
        ('38', 'SQ', 'Block Inc', 'Technology', '10M+', 'Payments/Bitcoin'),
        ('39', 'PYPL', 'PayPal Holdings', 'Financial', '15M+', 'Digital payments'),
        ('40', 'DIS', 'Walt Disney', 'Communication', '12M+', 'Entertainment'),
        ('41', 'NFLX', 'Netflix', 'Communication', '5M+', 'Streaming leader'),
        ('42', 'UBER', 'Uber Technologies', 'Technology', '20M+', 'Ride-share'),
        ('43', 'ABNB', 'Airbnb', 'Consumer Cyclical', '5M+', 'Travel platform'),
        ('44', 'CCL', 'Carnival Corp', 'Consumer Cyclical', '30M+', 'Cruise recovery'),
        ('45', 'AAL', 'American Airlines', 'Industrials', '30M+', 'Airline volatility'),
        ('46', 'DAL', 'Delta Air Lines', 'Industrials', '15M+', 'Airline leader'),
        ('47', 'UAL', 'United Airlines', 'Industrials', '10M+', 'Airline play'),
        ('48', 'BA', 'Boeing', 'Industrials', '8M+', 'Aerospace giant'),
        ('49', 'CAT', 'Caterpillar', 'Industrials', '3M+', 'Infrastructure'),
        ('50', 'DE', 'Deere & Co', 'Industrials', '2M+', 'Agriculture/AI'),
    ]

    for i, stock in enumerate(top_stocks):
        row = stocks_table.rows[i + 1].cells
        for j, val in enumerate(stock):
            row[j].text = val

    # ============ TOP 50 CRYPTOS ============
    doc.add_page_break()
    doc.add_heading('Top 50 Cryptos for Day Trading', level=1)
    doc.add_paragraph('Selected based on: 24h volume, volatility, market cap, and liquidity on major exchanges.')

    # Create crypto table
    crypto_table = doc.add_table(rows=51, cols=6)
    crypto_table.style = 'Table Grid'

    # Header
    headers = ['#', 'Symbol', 'Name', 'Category', '24h Volume', 'Why Selected']
    for i, header in enumerate(headers):
        crypto_table.rows[0].cells[i].text = header

    # Top 50 day trading cryptos
    top_cryptos = [
        ('1', 'BTC', 'Bitcoin', 'Store of Value', '$30B+', 'Most liquid crypto'),
        ('2', 'ETH', 'Ethereum', 'Smart Contract', '$15B+', 'DeFi backbone'),
        ('3', 'SOL', 'Solana', 'Smart Contract', '$3B+', 'Fast & cheap'),
        ('4', 'XRP', 'Ripple', 'Payments', '$2B+', 'Bank adoption'),
        ('5', 'DOGE', 'Dogecoin', 'Meme', '$1B+', 'Meme leader'),
        ('6', 'ADA', 'Cardano', 'Smart Contract', '$500M+', 'Academic approach'),
        ('7', 'AVAX', 'Avalanche', 'Smart Contract', '$400M+', 'DeFi growth'),
        ('8', 'SHIB', 'Shiba Inu', 'Meme', '$300M+', 'Meme volatility'),
        ('9', 'DOT', 'Polkadot', 'Infrastructure', '$300M+', 'Interoperability'),
        ('10', 'LINK', 'Chainlink', 'Oracle', '$400M+', 'Data provider'),
        ('11', 'MATIC', 'Polygon', 'Layer 2', '$400M+', 'Ethereum scaling'),
        ('12', 'UNI', 'Uniswap', 'DeFi', '$200M+', 'DEX leader'),
        ('13', 'LTC', 'Litecoin', 'Payments', '$500M+', 'Digital silver'),
        ('14', 'BCH', 'Bitcoin Cash', 'Payments', '$300M+', 'BTC fork'),
        ('15', 'ATOM', 'Cosmos', 'Infrastructure', '$200M+', 'Internet of blockchains'),
        ('16', 'FIL', 'Filecoin', 'Storage', '$200M+', 'Decentralized storage'),
        ('17', 'APT', 'Aptos', 'Smart Contract', '$200M+', 'Move language'),
        ('18', 'ARB', 'Arbitrum', 'Layer 2', '$300M+', 'ETH L2 leader'),
        ('19', 'OP', 'Optimism', 'Layer 2', '$150M+', 'ETH L2'),
        ('20', 'INJ', 'Injective', 'DeFi', '$200M+', 'DeFi derivatives'),
        ('21', 'SUI', 'Sui', 'Smart Contract', '$300M+', 'New L1'),
        ('22', 'SEI', 'Sei', 'Smart Contract', '$200M+', 'Trading focused'),
        ('23', 'NEAR', 'NEAR Protocol', 'Smart Contract', '$200M+', 'Scalable L1'),
        ('24', 'ICP', 'Internet Computer', 'Infrastructure', '$100M+', 'Web3 compute'),
        ('25', 'RENDER', 'Render', 'AI/GPU', '$150M+', 'GPU rendering'),
        ('26', 'FET', 'Fetch.ai', 'AI', '$200M+', 'AI agents'),
        ('27', 'AGIX', 'SingularityNET', 'AI', '$100M+', 'AI marketplace'),
        ('28', 'TAO', 'Bittensor', 'AI', '$100M+', 'Decentralized AI'),
        ('29', 'WLD', 'Worldcoin', 'Identity', '$150M+', 'Digital identity'),
        ('30', 'PEPE', 'Pepe', 'Meme', '$500M+', 'Top meme coin'),
        ('31', 'FLOKI', 'Floki Inu', 'Meme', '$100M+', 'Meme ecosystem'),
        ('32', 'BONK', 'Bonk', 'Meme', '$200M+', 'Solana meme'),
        ('33', 'WIF', 'dogwifhat', 'Meme', '$300M+', 'Solana meme'),
        ('34', 'AAVE', 'Aave', 'DeFi', '$100M+', 'Lending protocol'),
        ('35', 'MKR', 'Maker', 'DeFi', '$80M+', 'DAI stablecoin'),
        ('36', 'CRV', 'Curve', 'DeFi', '$100M+', 'Stablecoin DEX'),
        ('37', 'LDO', 'Lido DAO', 'DeFi', '$100M+', 'Liquid staking'),
        ('38', 'RUNE', 'THORChain', 'DeFi', '$100M+', 'Cross-chain DEX'),
        ('39', 'STX', 'Stacks', 'Bitcoin L2', '$100M+', 'Bitcoin DeFi'),
        ('40', 'IMX', 'Immutable', 'Gaming', '$80M+', 'NFT gaming'),
        ('41', 'SAND', 'The Sandbox', 'Metaverse', '$100M+', 'Virtual world'),
        ('42', 'MANA', 'Decentraland', 'Metaverse', '$80M+', 'Virtual real estate'),
        ('43', 'AXS', 'Axie Infinity', 'Gaming', '$50M+', 'Play-to-earn'),
        ('44', 'GALA', 'Gala Games', 'Gaming', '$100M+', 'Gaming platform'),
        ('45', 'ENS', 'Ethereum Name Service', 'Infrastructure', '$50M+', 'Domain names'),
        ('46', 'GRT', 'The Graph', 'Infrastructure', '$80M+', 'Indexing protocol'),
        ('47', 'ALGO', 'Algorand', 'Smart Contract', '$50M+', 'Pure PoS'),
        ('48', 'XLM', 'Stellar', 'Payments', '$100M+', 'Cross-border'),
        ('49', 'HBAR', 'Hedera', 'Infrastructure', '$80M+', 'Enterprise grade'),
        ('50', 'VET', 'VeChain', 'Supply Chain', '$100M+', 'Enterprise blockchain'),
    ]

    for i, crypto in enumerate(top_cryptos):
        row = crypto_table.rows[i + 1].cells
        for j, val in enumerate(crypto):
            row[j].text = val

    # Trading Tips
    doc.add_page_break()
    doc.add_heading('Day Trading Tips', level=1)

    tips = [
        'Focus on assets with average daily volume > $10M for stocks, > $100M for cryptos',
        'Look for volatility (ATR) of 2-5% for optimal day trading opportunities',
        'Trade during peak hours: 9:30-11:30 AM ET for stocks, 24/7 for crypto with higher volume during US hours',
        'Use technical indicators: RSI for overbought/oversold, MACD for momentum',
        'Set stop-losses at 1-2% to manage risk',
        'The "Active Trading List" in the app is auto-generated weekly based on these criteria'
    ]

    for tip in tips:
        doc.add_paragraph(tip, style='List Bullet')

    # Save
    doc.save('TOP_100_STOCKS_CRYPTOS.docx')
    print("Created: TOP_100_STOCKS_CRYPTOS.docx")

if __name__ == "__main__":
    create_top_100_doc()
