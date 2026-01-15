"""
Create Word Document with Weekly Table Structures
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_weekly_table_doc():
    doc = Document()

    # Title
    title = doc.add_heading('Weekly Data Table Structures', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('BigQuery Tables for Weekly Stock and Crypto Analysis')
    doc.add_paragraph(f'Project: cryptobot-462709')
    doc.add_paragraph(f'Dataset: crypto_trading_data')

    # ============ STOCKS WEEKLY SUMMARY ============
    doc.add_heading('1. stocks_weekly_summary (58 columns)', level=1)
    doc.add_paragraph('Weekly summary for ALL US stocks (~20,000). Updated every Saturday at 4:00 AM ET.')

    # Basic Information
    doc.add_heading('Basic Information', level=2)
    basic_fields = [
        ('symbol', 'STRING', 'REQUIRED', 'Stock ticker symbol (e.g., AAPL, MSFT)'),
        ('name', 'STRING', 'NULLABLE', 'Company name'),
        ('exchange', 'STRING', 'NULLABLE', 'Exchange (NYSE, NASDAQ, AMEX)'),
        ('type', 'STRING', 'NULLABLE', 'Security type (Common Stock, ETF)'),
        ('sector', 'STRING', 'NULLABLE', 'Business sector (Technology, Healthcare, etc.)'),
        ('industry', 'STRING', 'NULLABLE', 'Industry within sector'),
        ('country', 'STRING', 'NULLABLE', 'Country of listing'),
        ('currency', 'STRING', 'NULLABLE', 'Trading currency (USD)'),
    ]
    add_field_table(doc, basic_fields)

    # Price Data
    doc.add_heading('Price Data', level=2)
    price_fields = [
        ('current_price', 'FLOAT64', 'NULLABLE', 'Current/Latest stock price'),
        ('open_price', 'FLOAT64', 'NULLABLE', 'Week opening price'),
        ('high_price', 'FLOAT64', 'NULLABLE', 'Week high price'),
        ('low_price', 'FLOAT64', 'NULLABLE', 'Week low price'),
        ('close_price', 'FLOAT64', 'NULLABLE', 'Week closing price'),
        ('previous_close', 'FLOAT64', 'NULLABLE', 'Previous week closing price'),
    ]
    add_field_table(doc, price_fields)

    # Volume Data
    doc.add_heading('Volume Data', level=2)
    volume_fields = [
        ('volume', 'INT64', 'NULLABLE', 'Weekly trading volume'),
        ('average_volume', 'INT64', 'NULLABLE', 'Average daily volume'),
    ]
    add_field_table(doc, volume_fields)

    # Performance Metrics
    doc.add_heading('Performance Metrics', level=2)
    perf_fields = [
        ('weekly_change', 'FLOAT64', 'NULLABLE', 'Price change this week ($)'),
        ('weekly_change_percent', 'FLOAT64', 'NULLABLE', 'Price change this week (%)'),
        ('monthly_change_percent', 'FLOAT64', 'NULLABLE', '30-day price change (%)'),
        ('ytd_change_percent', 'FLOAT64', 'NULLABLE', 'Year-to-date change (%)'),
    ]
    add_field_table(doc, perf_fields)

    # Valuation Metrics
    doc.add_heading('Valuation Metrics', level=2)
    val_fields = [
        ('market_cap', 'FLOAT64', 'NULLABLE', 'Market capitalization in USD'),
        ('pe_ratio', 'FLOAT64', 'NULLABLE', 'Price-to-Earnings ratio'),
        ('eps', 'FLOAT64', 'NULLABLE', 'Earnings per share'),
        ('dividend_yield', 'FLOAT64', 'NULLABLE', 'Dividend yield (%)'),
        ('beta', 'FLOAT64', 'NULLABLE', 'Beta (volatility vs market)'),
    ]
    add_field_table(doc, val_fields)

    # Trading Ranges
    doc.add_heading('Trading Ranges', level=2)
    range_fields = [
        ('week_52_high', 'FLOAT64', 'NULLABLE', '52-week high price'),
        ('week_52_low', 'FLOAT64', 'NULLABLE', '52-week low price'),
        ('percent_from_52_high', 'FLOAT64', 'NULLABLE', '% below 52-week high'),
        ('percent_from_52_low', 'FLOAT64', 'NULLABLE', '% above 52-week low'),
    ]
    add_field_table(doc, range_fields)

    # Volatility Metrics
    doc.add_heading('Volatility Metrics (Day Trading)', level=2)
    vol_fields = [
        ('volatility_weekly', 'FLOAT64', 'NULLABLE', 'Weekly price volatility %'),
        ('volatility_monthly', 'FLOAT64', 'NULLABLE', 'Monthly price volatility %'),
        ('atr', 'FLOAT64', 'NULLABLE', 'Average True Range'),
        ('atr_percent', 'FLOAT64', 'NULLABLE', 'ATR as % of price'),
    ]
    add_field_table(doc, vol_fields)

    # Day Trading Scores
    doc.add_heading('Day Trading Scores', level=2)
    score_fields = [
        ('day_trade_score', 'FLOAT64', 'NULLABLE', 'Day trading suitability score (0-100)'),
        ('liquidity_score', 'FLOAT64', 'NULLABLE', 'Liquidity score (0-100)'),
        ('momentum_score', 'FLOAT64', 'NULLABLE', 'Momentum score (0-100)'),
    ]
    add_field_table(doc, score_fields)

    # Technical Indicators
    doc.add_heading('Technical Indicators', level=2)
    tech_fields = [
        ('rsi_weekly', 'FLOAT64', 'NULLABLE', 'Weekly RSI (Relative Strength Index)'),
        ('macd_weekly', 'FLOAT64', 'NULLABLE', 'Weekly MACD'),
        ('macd_signal_weekly', 'FLOAT64', 'NULLABLE', 'Weekly MACD Signal line'),
        ('sma_20', 'FLOAT64', 'NULLABLE', '20-day Simple Moving Average'),
        ('sma_50', 'FLOAT64', 'NULLABLE', '50-day Simple Moving Average'),
        ('sma_200', 'FLOAT64', 'NULLABLE', '200-day Simple Moving Average'),
    ]
    add_field_table(doc, tech_fields)

    # Trend Analysis
    doc.add_heading('Trend Analysis', level=2)
    trend_fields = [
        ('trend_short', 'STRING', 'NULLABLE', 'Short-term trend (bullish/bearish/neutral)'),
        ('trend_medium', 'STRING', 'NULLABLE', 'Medium-term trend'),
        ('trend_long', 'STRING', 'NULLABLE', 'Long-term trend'),
        ('above_sma_20', 'BOOLEAN', 'NULLABLE', 'Price above 20-day SMA'),
        ('above_sma_50', 'BOOLEAN', 'NULLABLE', 'Price above 50-day SMA'),
        ('above_sma_200', 'BOOLEAN', 'NULLABLE', 'Price above 200-day SMA'),
    ]
    add_field_table(doc, trend_fields)

    # Classifications (NLP Search)
    doc.add_heading('Classifications (for NLP Search)', level=2)
    class_fields = [
        ('market_cap_category', 'STRING', 'NULLABLE', 'mega/large/mid/small/micro cap'),
        ('volatility_category', 'STRING', 'NULLABLE', 'high/medium/low volatility'),
        ('momentum_category', 'STRING', 'NULLABLE', 'strong_up/up/neutral/down/strong_down'),
        ('value_category', 'STRING', 'NULLABLE', 'undervalued/fair/overvalued'),
    ]
    add_field_table(doc, class_fields)

    # Metadata
    doc.add_heading('Metadata', level=2)
    meta_fields = [
        ('week_start_date', 'DATE', 'NULLABLE', 'Start of the week'),
        ('week_end_date', 'DATE', 'NULLABLE', 'End of the week'),
        ('fetch_timestamp', 'TIMESTAMP', 'REQUIRED', 'When data was fetched'),
        ('data_source', 'STRING', 'NULLABLE', 'Data source (twelvedata)'),
    ]
    add_field_table(doc, meta_fields)

    # Active List
    doc.add_heading('Active List Flags', level=2)
    active_fields = [
        ('is_active_pick', 'BOOLEAN', 'NULLABLE', 'Selected for active trading list'),
        ('active_pick_reason', 'STRING', 'NULLABLE', 'Why selected for active list'),
    ]
    add_field_table(doc, active_fields)

    # ============ CRYPTOS WEEKLY SUMMARY ============
    doc.add_page_break()
    doc.add_heading('2. cryptos_weekly_summary (55 columns)', level=1)
    doc.add_paragraph('Weekly summary for ALL cryptos. Updated every Saturday at 4:30 AM ET.')

    # Basic Information
    doc.add_heading('Basic Information', level=2)
    crypto_basic = [
        ('symbol', 'STRING', 'REQUIRED', 'Crypto symbol (BTC, ETH, etc.)'),
        ('name', 'STRING', 'NULLABLE', 'Full name (Bitcoin, Ethereum, etc.)'),
        ('pair', 'STRING', 'NULLABLE', 'Trading pair (BTC/USD)'),
        ('category', 'STRING', 'NULLABLE', 'Category (DeFi, Layer1, Meme, Layer2, AI, etc.)'),
        ('subcategory', 'STRING', 'NULLABLE', 'Subcategory'),
    ]
    add_field_table(doc, crypto_basic)

    # Price Data
    doc.add_heading('Price Data', level=2)
    crypto_price = [
        ('current_price', 'FLOAT64', 'NULLABLE', 'Current/Latest price'),
        ('open_price', 'FLOAT64', 'NULLABLE', 'Week opening price'),
        ('high_price', 'FLOAT64', 'NULLABLE', 'Week high price'),
        ('low_price', 'FLOAT64', 'NULLABLE', 'Week low price'),
        ('close_price', 'FLOAT64', 'NULLABLE', 'Week closing price'),
        ('previous_close', 'FLOAT64', 'NULLABLE', 'Previous week closing price'),
    ]
    add_field_table(doc, crypto_price)

    # Volume Data
    doc.add_heading('Volume Data', level=2)
    crypto_vol = [
        ('volume_24h', 'FLOAT64', 'NULLABLE', '24-hour trading volume USD'),
        ('volume_weekly', 'FLOAT64', 'NULLABLE', 'Weekly trading volume USD'),
    ]
    add_field_table(doc, crypto_vol)

    # Performance Metrics
    doc.add_heading('Performance Metrics', level=2)
    crypto_perf = [
        ('weekly_change', 'FLOAT64', 'NULLABLE', 'Price change this week ($)'),
        ('weekly_change_percent', 'FLOAT64', 'NULLABLE', 'Price change this week (%)'),
        ('monthly_change_percent', 'FLOAT64', 'NULLABLE', '30-day price change (%)'),
        ('ytd_change_percent', 'FLOAT64', 'NULLABLE', 'Year-to-date change (%)'),
    ]
    add_field_table(doc, crypto_perf)

    # Market Data
    doc.add_heading('Market Data', level=2)
    crypto_market = [
        ('market_cap', 'FLOAT64', 'NULLABLE', 'Market capitalization USD'),
        ('market_cap_rank', 'INT64', 'NULLABLE', 'Market cap ranking'),
        ('circulating_supply', 'FLOAT64', 'NULLABLE', 'Circulating supply'),
        ('total_supply', 'FLOAT64', 'NULLABLE', 'Total supply'),
        ('max_supply', 'FLOAT64', 'NULLABLE', 'Maximum supply'),
    ]
    add_field_table(doc, crypto_market)

    # All Time Highs/Lows
    doc.add_heading('All-Time High/Low', level=2)
    crypto_ath = [
        ('ath', 'FLOAT64', 'NULLABLE', 'All-time high price'),
        ('ath_date', 'DATE', 'NULLABLE', 'Date of all-time high'),
        ('percent_from_ath', 'FLOAT64', 'NULLABLE', '% below all-time high'),
        ('atl', 'FLOAT64', 'NULLABLE', 'All-time low price'),
        ('atl_date', 'DATE', 'NULLABLE', 'Date of all-time low'),
    ]
    add_field_table(doc, crypto_ath)

    # Same structure continues for volatility, scores, indicators, trends, classifications, metadata, active list

    # ============ ACTIVE TRADING LIST ============
    doc.add_page_break()
    doc.add_heading('3. active_trading_list (16 columns)', level=1)
    doc.add_paragraph('Top 100 picks for day trading. Generated from weekly analysis.')

    active_list_fields = [
        ('asset_type', 'STRING', 'REQUIRED', 'stock or crypto'),
        ('symbol', 'STRING', 'REQUIRED', 'Ticker/Symbol'),
        ('name', 'STRING', 'NULLABLE', 'Company/Crypto name'),
        ('current_price', 'FLOAT64', 'NULLABLE', 'Current price'),
        ('sector', 'STRING', 'NULLABLE', 'Sector/Category'),
        ('selection_rank', 'INT64', 'NULLABLE', 'Ranking 1-100'),
        ('day_trade_score', 'FLOAT64', 'NULLABLE', 'Day trading score (0-100)'),
        ('selection_reason', 'STRING', 'NULLABLE', 'Why selected'),
        ('weekly_change_percent', 'FLOAT64', 'NULLABLE', 'Weekly change %'),
        ('volatility_percent', 'FLOAT64', 'NULLABLE', 'Volatility %'),
        ('volume_ratio', 'FLOAT64', 'NULLABLE', 'Volume vs average'),
        ('rsi', 'FLOAT64', 'NULLABLE', 'RSI'),
        ('trend', 'STRING', 'NULLABLE', 'Overall trend'),
        ('generated_date', 'DATE', 'REQUIRED', 'When list was generated'),
        ('valid_until', 'DATE', 'NULLABLE', 'List valid until'),
        ('fetch_timestamp', 'TIMESTAMP', 'REQUIRED', 'Timestamp'),
    ]
    add_field_table(doc, active_list_fields)

    # ============ SCHEDULER EXECUTION LOG ============
    doc.add_page_break()
    doc.add_heading('4. scheduler_execution_log (30 columns)', level=1)
    doc.add_paragraph('Tracks all scheduler runs for Admin monitoring.')

    scheduler_fields = [
        ('execution_id', 'STRING', 'REQUIRED', 'Unique execution ID'),
        ('scheduler_name', 'STRING', 'REQUIRED', 'Name of the scheduler job'),
        ('function_name', 'STRING', 'REQUIRED', 'Cloud function name'),
        ('table_name', 'STRING', 'REQUIRED', 'Target BigQuery table'),
        ('execution_date', 'DATE', 'REQUIRED', 'Date of execution'),
        ('start_time', 'TIMESTAMP', 'REQUIRED', 'When execution started'),
        ('end_time', 'TIMESTAMP', 'NULLABLE', 'When execution ended'),
        ('duration_seconds', 'FLOAT64', 'NULLABLE', 'Total execution time in seconds'),
        ('duration_minutes', 'FLOAT64', 'NULLABLE', 'Total execution time in minutes'),
        ('status', 'STRING', 'REQUIRED', 'SUCCESS, FAILED, PARTIAL, RUNNING'),
        ('error_message', 'STRING', 'NULLABLE', 'Error details if failed'),
        ('total_symbols', 'INT64', 'NULLABLE', 'Total symbols to process'),
        ('successful_symbols', 'INT64', 'NULLABLE', 'Successfully processed'),
        ('failed_symbols', 'INT64', 'NULLABLE', 'Failed to process'),
        ('records_inserted', 'INT64', 'NULLABLE', 'Records inserted into table'),
        ('api_calls_made', 'INT64', 'NULLABLE', 'Number of API calls'),
        ('was_manual_trigger', 'BOOLEAN', 'NULLABLE', 'Manually triggered vs scheduled'),
        ('created_at', 'TIMESTAMP', 'REQUIRED', 'Log entry creation time'),
    ]
    add_field_table(doc, scheduler_fields)

    # Save
    doc.save('WEEKLY_TABLE_STRUCTURES.docx')
    print("Created: WEEKLY_TABLE_STRUCTURES.docx")


def add_field_table(doc, fields):
    """Add a table with field definitions"""
    table = doc.add_table(rows=len(fields) + 1, cols=4)
    table.style = 'Table Grid'

    # Header row
    headers = ['Field Name', 'Type', 'Mode', 'Description']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header

    # Data rows
    for i, (name, dtype, mode, desc) in enumerate(fields):
        row = table.rows[i + 1].cells
        row[0].text = name
        row[1].text = dtype
        row[2].text = mode
        row[3].text = desc

    doc.add_paragraph('')  # Add spacing


if __name__ == "__main__":
    create_weekly_table_doc()
