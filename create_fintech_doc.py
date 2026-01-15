from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('AIAlgoTradeHits.com - Fintech AI Data Architecture & Implementation Specification', 0)

# Subtitle info
doc.add_paragraph('Master Technical Document')
info = doc.add_paragraph()
info.add_run('Version: ').bold = True
info.add_run('2.0\n')
info.add_run('Platform: ').bold = True
info.add_run('Google Cloud Platform (GCP)\n')
info.add_run('AI Engine: ').bold = True
info.add_run('Vertex AI + Gemini 3\n')
info.add_run('Date: ').bold = True
info.add_run('November 2025\n')
info.add_run('Classification: ').bold = True
info.add_run('Technical Specification')

doc.add_paragraph()

# Executive Summary
doc.add_heading('Executive Summary', level=1)
doc.add_paragraph('AIAlgoTradeHits.com is a comprehensive cross-asset, AI-driven analytics, alerting, and intelligence platform delivering institutional-grade insights across equities, ETFs, cryptocurrency, forex, commodities, and interest-rate markets. Built entirely on Google Cloud Platform, the system leverages Vertex AI, Gemini 3 LLM, BigQuery, and Cloud Run to create a next-generation fintech trading intelligence engine.')

doc.add_heading('Platform Vision', level=2)
table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Component'
hdr[1].text = 'Technology'
hdr[2].text = 'Purpose'
data = [
    ('Data Warehouse', 'BigQuery', 'Multi-layer data architecture (Bronze/Silver/Gold)'),
    ('AI/ML Engine', 'Vertex AI', 'Predictive models, pattern recognition, regime classification'),
    ('LLM Intelligence', 'Gemini 3 Pro', 'Natural language analysis, RAG-powered commentary'),
    ('Document Intelligence', 'Document AI + RAG', 'Research document processing, knowledge retrieval'),
    ('Real-Time Serving', 'Cloud Run + Cloud Functions', 'API endpoints, scheduled data collection'),
    ('Frontend', 'React + TradingView Charts', 'Professional trading interface'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

doc.add_paragraph()

# Part 1: Data Architecture
doc.add_heading('Part 1: Data Architecture', level=1)

doc.add_heading('1.1 Multi-Layer Data Architecture', level=2)
doc.add_paragraph('The platform implements a rigorous four-layer data architecture inspired by institutional quant trading desks:')
doc.add_paragraph('BRONZE LAYER - Raw, Immutable Market Data: Exact API responses preserved, full lineage and audit trail, compliance-grade immutability')
doc.add_paragraph('SILVER LAYER - Cleaned & Standardized Data: Normalized timestamps (UTC), missing data handling, anomaly removal, session calendar alignment')
doc.add_paragraph('GOLD LAYER - Engineered Features & Intelligence: 70+ Technical Indicators, pattern recognition signals, cross-asset correlations, volatility regime classification, ML-ready feature vectors')
doc.add_paragraph('WIDE TABLES - Materialized Fast-Access Layer: Flattened serving tables, real-time UI/API access, LLM RAG retrieval optimized, ML inference ready')

doc.add_heading('1.2 BigQuery Table Structure', level=2)
doc.add_heading('1.2.1 Asset Class Tables (6 Categories)', level=3)
table = doc.add_table(rows=9, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Table Name'
hdr[1].text = 'Asset Class'
hdr[2].text = 'Timeframes'
hdr[3].text = 'Data Source'
data = [
    ('stocks_historical_daily', 'US Equities', 'Daily, Weekly', 'TwelveData'),
    ('stocks_hourly', 'US Equities', 'Hourly, 5-min', 'TwelveData'),
    ('cryptos_historical_daily', 'Cryptocurrency', 'Daily, Weekly', 'Kraken/TwelveData'),
    ('crypto_hourly_data', 'Cryptocurrency', 'Hourly, 5-min', 'Kraken'),
    ('etfs_historical_daily', 'ETFs', 'Daily, Weekly', 'TwelveData'),
    ('forex_historical_daily', 'FX Pairs', 'Daily, Weekly', 'TwelveData'),
    ('indices_historical_daily', 'Market Indices', 'Daily, Weekly', 'TwelveData'),
    ('commodities_historical_daily', 'Commodities', 'Daily, Weekly', 'TwelveData'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]
    row[3].text = row_data[3]

doc.add_paragraph()

doc.add_heading('1.2.2 Complete Field Schema (Phase 1 Enhanced)', level=3)
doc.add_paragraph('Based on Phase 1 Methodology requirements, the following fields are required for ML training:')

# Core OHLCV Fields
doc.add_heading('Core OHLCV Fields (Required)', level=3)
table = doc.add_table(rows=10, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Field Name'
hdr[1].text = 'Type'
hdr[2].text = 'Description'
hdr[3].text = 'Status'
data = [
    ('symbol', 'STRING', 'Asset ticker symbol', 'Available'),
    ('datetime', 'TIMESTAMP', 'Candle timestamp (UTC)', 'Available'),
    ('open', 'FLOAT64', 'Opening price', 'Available'),
    ('high', 'FLOAT64', 'High price', 'Available'),
    ('low', 'FLOAT64', 'Low price', 'Available'),
    ('close', 'FLOAT64', 'Closing price', 'Available'),
    ('volume', 'FLOAT64', 'Trading volume', 'Available'),
    ('candle_body_pct', 'FLOAT64', '(close - open) / open', 'Add'),
    ('candle_range_pct', 'FLOAT64', '(high - low) / open', 'Add'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]
    row[3].text = row_data[3]

doc.add_paragraph()

# Return & Momentum Fields
doc.add_heading('Return & Momentum Fields', level=3)
table = doc.add_table(rows=6, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Field Name'
hdr[1].text = 'Type'
hdr[2].text = 'Description'
hdr[3].text = 'Status'
data = [
    ('weekly_return', 'FLOAT64', 'Weekly percentage return', 'Available'),
    ('weekly_log_return', 'FLOAT64', 'ln(close_t / close_t-1)', 'Add'),
    ('return_2w', 'FLOAT64', '2-week return', 'Add'),
    ('return_4w', 'FLOAT64', '4-week return', 'Add'),
    ('daily_return_pct', 'FLOAT64', 'Daily return percentage', 'Available'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]
    row[3].text = row_data[3]

doc.add_paragraph()

# RSI Indicators
doc.add_heading('RSI Indicators', level=3)
table = doc.add_table(rows=8, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Field Name'
hdr[1].text = 'Type'
hdr[2].text = 'Description'
hdr[3].text = 'Status'
data = [
    ('rsi_7', 'FLOAT64', 'RSI 7-period', 'Available'),
    ('rsi_14', 'FLOAT64', 'RSI 14-period', 'Available'),
    ('rsi_21', 'FLOAT64', 'RSI 21-period', 'Available'),
    ('rsi_slope', 'FLOAT64', 'RSI change rate', 'Add'),
    ('rsi_zscore', 'FLOAT64', 'RSI z-score (100-week)', 'Add'),
    ('rsi_overbought_flag', 'INT64', 'RSI > 70 flag', 'Add'),
    ('rsi_oversold_flag', 'INT64', 'RSI < 30 flag', 'Add'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]
    row[3].text = row_data[3]

doc.add_paragraph()

# MACD Indicators
doc.add_heading('MACD Indicators', level=3)
table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Field Name'
hdr[1].text = 'Type'
hdr[2].text = 'Description'
hdr[3].text = 'Status'
data = [
    ('macd', 'FLOAT64', 'MACD line (12,26)', 'Available'),
    ('macd_signal', 'FLOAT64', 'MACD signal (9)', 'Available'),
    ('macd_hist', 'FLOAT64', 'MACD histogram', 'Add'),
    ('macd_cross_flag', 'INT64', 'Cross signal (+1/-1/0)', 'Add'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]
    row[3].text = row_data[3]

doc.add_paragraph()

# Moving Averages - SMA
doc.add_heading('Moving Averages - SMA', level=3)
table = doc.add_table(rows=7, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Field Name'
hdr[1].text = 'Type'
hdr[2].text = 'Description'
hdr[3].text = 'Status'
data = [
    ('sma_5', 'FLOAT64', 'SMA 5-period', 'Available'),
    ('sma_10', 'FLOAT64', 'SMA 10-period', 'Available'),
    ('sma_20', 'FLOAT64', 'SMA 20-period', 'Available'),
    ('sma_50', 'FLOAT64', 'SMA 50-period', 'Available'),
    ('sma_100', 'FLOAT64', 'SMA 100-period', 'Available'),
    ('sma_200', 'FLOAT64', 'SMA 200-period', 'Available'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]
    row[3].text = row_data[3]

doc.add_paragraph()

# Moving Averages - EMA
doc.add_heading('Moving Averages - EMA', level=3)
table = doc.add_table(rows=7, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Field Name'
hdr[1].text = 'Type'
hdr[2].text = 'Description'
hdr[3].text = 'Status'
data = [
    ('ema_5', 'FLOAT64', 'EMA 5-period', 'Add'),
    ('ema_10', 'FLOAT64', 'EMA 10-period', 'Add'),
    ('ema_20', 'FLOAT64', 'EMA 20-period', 'Add'),
    ('ema_50', 'FLOAT64', 'EMA 50-period', 'Add'),
    ('ema_100', 'FLOAT64', 'EMA 100-period', 'Add'),
    ('ema_200', 'FLOAT64', 'EMA 200-period', 'Add'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]
    row[3].text = row_data[3]

doc.add_paragraph()

# MA Distance Features
doc.add_heading('MA Distance Features', level=3)
table = doc.add_table(rows=7, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Field Name'
hdr[1].text = 'Type'
hdr[2].text = 'Description'
hdr[3].text = 'Status'
data = [
    ('close_vs_sma20_pct', 'FLOAT64', '(close / SMA20 - 1) * 100', 'Add'),
    ('close_vs_sma50_pct', 'FLOAT64', '(close / SMA50 - 1) * 100', 'Add'),
    ('close_vs_sma200_pct', 'FLOAT64', '(close / SMA200 - 1) * 100', 'Add'),
    ('close_vs_ema20_pct', 'FLOAT64', '(close / EMA20 - 1) * 100', 'Add'),
    ('ema_slope_20', 'FLOAT64', 'EMA20 slope', 'Add'),
    ('ema_slope_50', 'FLOAT64', 'EMA50 slope', 'Add'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]
    row[3].text = row_data[3]

doc.add_paragraph()

# Volatility Indicators
doc.add_heading('Volatility Indicators', level=3)
table = doc.add_table(rows=8, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Field Name'
hdr[1].text = 'Type'
hdr[2].text = 'Description'
hdr[3].text = 'Status'
data = [
    ('atr_7', 'FLOAT64', 'ATR 7-period', 'Available'),
    ('atr_14', 'FLOAT64', 'ATR 14-period', 'Available'),
    ('atr_pct', 'FLOAT64', 'ATR / close * 100', 'Add'),
    ('atr_zscore', 'FLOAT64', 'ATR z-score', 'Add'),
    ('atr_slope', 'FLOAT64', 'ATR change rate', 'Add'),
    ('volatility_10', 'FLOAT64', '10-period volatility', 'Available'),
    ('volatility_20', 'FLOAT64', '20-period volatility', 'Available'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]
    row[3].text = row_data[3]

doc.add_paragraph()

# Bollinger Bands
doc.add_heading('Bollinger Bands', level=3)
table = doc.add_table(rows=6, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Field Name'
hdr[1].text = 'Type'
hdr[2].text = 'Description'
hdr[3].text = 'Status'
data = [
    ('bb_upper', 'FLOAT64', 'Upper band (20, 2 sigma)', 'Available'),
    ('bb_middle', 'FLOAT64', 'Middle band (SMA20)', 'Available'),
    ('bb_lower', 'FLOAT64', 'Lower band (20, 2 sigma)', 'Available'),
    ('bb_width', 'FLOAT64', 'Band width percentage', 'Available'),
    ('bb_percent', 'FLOAT64', '%B indicator', 'Available'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]
    row[3].text = row_data[3]

doc.add_paragraph()

# Stochastic & Oscillators
doc.add_heading('Stochastic & Oscillators', level=3)
table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Field Name'
hdr[1].text = 'Type'
hdr[2].text = 'Description'
hdr[3].text = 'Status'
data = [
    ('stoch_k', 'FLOAT64', 'Stochastic %K', 'Available'),
    ('stoch_d', 'FLOAT64', 'Stochastic %D', 'Available'),
    ('williams_r', 'FLOAT64', 'Williams %R', 'Available'),
    ('cci', 'FLOAT64', 'CCI (20)', 'Available'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]
    row[3].text = row_data[3]

doc.add_paragraph()

# Volume Indicators
doc.add_heading('Volume Indicators', level=3)
table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Field Name'
hdr[1].text = 'Type'
hdr[2].text = 'Description'
hdr[3].text = 'Status'
data = [
    ('obv', 'FLOAT64', 'On-Balance Volume', 'Available'),
    ('volume_zscore', 'FLOAT64', 'Volume z-score', 'Add'),
    ('volume_ratio', 'FLOAT64', 'Volume / MA volume', 'Add'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]
    row[3].text = row_data[3]

doc.add_paragraph()

# Trend Indicators (ADX)
doc.add_heading('Trend Indicators (ADX)', level=3)
table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Field Name'
hdr[1].text = 'Type'
hdr[2].text = 'Description'
hdr[3].text = 'Status'
data = [
    ('adx', 'FLOAT64', 'ADX (14)', 'Add'),
    ('plus_di', 'FLOAT64', '+DI (14)', 'Add'),
    ('minus_di', 'FLOAT64', '-DI (14)', 'Add'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]
    row[3].text = row_data[3]

doc.add_paragraph()

# Pivot & Structure Features
doc.add_heading('Pivot & Structure Features', level=3)
table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Field Name'
hdr[1].text = 'Type'
hdr[2].text = 'Description'
hdr[3].text = 'Status'
data = [
    ('pivot_high_flag', 'INT64', 'Pivot high detected', 'Add'),
    ('pivot_low_flag', 'INT64', 'Pivot low detected', 'Add'),
    ('dist_to_pivot_high', 'FLOAT64', 'Distance to last pivot high', 'Add'),
    ('dist_to_pivot_low', 'FLOAT64', 'Distance to last pivot low', 'Add'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]
    row[3].text = row_data[3]

doc.add_paragraph()

# Regime & Classification
doc.add_heading('Regime & Classification', level=3)
table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Field Name'
hdr[1].text = 'Type'
hdr[2].text = 'Description'
hdr[3].text = 'Status'
data = [
    ('regime_state', 'INT64', 'Numeric regime (1-6)', 'Add'),
    ('regime_confidence', 'FLOAT64', 'Regime confidence (0-1)', 'Add'),
    ('momentum_signal', 'STRING', 'overbought/oversold/bullish/bearish', 'Available'),
    ('trend_signal', 'STRING', 'strong_bullish/bullish/bearish/strong_bearish', 'Available'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]
    row[3].text = row_data[3]

doc.add_paragraph()

# ML Target Variables
doc.add_heading('ML Target Variables', level=3)
table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Field Name'
hdr[1].text = 'Type'
hdr[2].text = 'Description'
hdr[3].text = 'Status'
data = [
    ('target_return_1d', 'FLOAT64', 'Next-day return %', 'Available'),
    ('target_return_5d', 'FLOAT64', '5-day forward return %', 'Available'),
    ('target_direction_1d', 'INT64', 'Next-day direction (0/1)', 'Available'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]
    row[3].text = row_data[3]

doc.add_paragraph()

# Lagged Features
doc.add_heading('Lagged Features', level=3)
table = doc.add_table(rows=9, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Field Name'
hdr[1].text = 'Type'
hdr[2].text = 'Description'
hdr[3].text = 'Status'
data = [
    ('close_lag_1', 'FLOAT64', 'Close t-1', 'Available'),
    ('close_lag_2', 'FLOAT64', 'Close t-2', 'Available'),
    ('close_lag_3', 'FLOAT64', 'Close t-3', 'Available'),
    ('close_lag_5', 'FLOAT64', 'Close t-5', 'Available'),
    ('close_lag_10', 'FLOAT64', 'Close t-10', 'Available'),
    ('return_lag_1', 'FLOAT64', 'Return t-1', 'Available'),
    ('return_lag_2', 'FLOAT64', 'Return t-2', 'Available'),
    ('return_lag_3', 'FLOAT64', 'Return t-3', 'Available'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]
    row[3].text = row_data[3]

doc.add_paragraph()

# Metadata Fields
doc.add_heading('Metadata Fields', level=3)
table = doc.add_table(rows=3, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Field Name'
hdr[1].text = 'Type'
hdr[2].text = 'Description'
hdr[3].text = 'Status'
data = [
    ('asset_type', 'STRING', 'stocks/cryptos/etfs/forex/indices/commodities', 'Available'),
    ('fetch_timestamp', 'TIMESTAMP', 'Data collection timestamp', 'Available'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]
    row[3].text = row_data[3]

doc.add_paragraph()

# Part 2: AI & Machine Learning Architecture
doc.add_heading('Part 2: AI & Machine Learning Architecture', level=1)

doc.add_heading('2.1 Vertex AI Integration', level=2)
doc.add_heading('2.1.1 Model Types & Use Cases', level=3)
table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Model Type'
hdr[1].text = 'Vertex AI Service'
hdr[2].text = 'Use Case'
data = [
    ('Gradient Boosting', 'AutoML Tables', 'Regime classification, direction prediction'),
    ('Time Series', 'Vertex AI Forecasting', 'Price prediction, volatility forecasting'),
    ('LSTM/Transformer', 'Custom Training', 'Sequential pattern learning'),
    ('Vision Models', 'AutoML Vision', 'Chart pattern recognition'),
    ('LLM', 'Gemini 3 Pro', 'Commentary generation, analysis'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

doc.add_paragraph()

doc.add_heading('2.1.2 ML Pipeline Architecture', level=3)
doc.add_paragraph('VERTEX AI PIPELINES consists of three main components:')
doc.add_paragraph('1. TRAINING PIPELINE: Feature extraction, model training, hyperparameter tuning')
doc.add_paragraph('2. PREDICTION PIPELINE: Real-time inference, batch predictions, alert generation')
doc.add_paragraph('3. EVALUATION PIPELINE: Backtesting, performance metrics, drift detection')

doc.add_paragraph()

doc.add_heading('2.2 Gemini 3 LLM Integration', level=2)
doc.add_heading('2.2.1 LLM Use Cases', level=3)
table = doc.add_table(rows=6, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Use Case'
hdr[1].text = 'Input'
hdr[2].text = 'Output'
hdr[3].text = 'Frequency'
data = [
    ('Market Commentary', 'Latest indicators, patterns, regime', 'Research-style analysis', 'Real-time'),
    ('Alert Explanations', 'Alert trigger data, context', 'Natural language reasoning', 'Per alert'),
    ('Pattern Analysis', 'Chart patterns, structure', 'Professional interpretation', 'On detection'),
    ('Risk Assessment', 'Cross-asset correlations', 'Risk narrative', 'Hourly'),
    ('Q&A Interface', 'User questions + RAG context', 'Informed responses', 'On-demand'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]
    row[3].text = row_data[3]

doc.add_paragraph()

doc.add_heading('2.2.2 RAG (Retrieval-Augmented Generation) Architecture', level=3)
doc.add_paragraph('The RAG pipeline consists of:')
doc.add_paragraph('1. USER QUERY / TRIGGER - Initial request or automated trigger')
doc.add_paragraph('2. CONTEXT RETRIEVAL - BigQuery Market Data, Vector DB Research Embeddings, Document Store (PDF/DOCX)')
doc.add_paragraph('3. PROMPT CONSTRUCTION - System Instructions + Retrieved Context + User Query')
doc.add_paragraph('4. GEMINI 3 PRO - Grounded response generation, citation of data sources, confidence scoring, hallucination prevention')
doc.add_paragraph('5. STRUCTURED OUTPUT - Market analysis narrative, key insights with citations, risk warnings, actionable recommendations')

doc.add_paragraph()

doc.add_heading('2.2.3 Document AI & Knowledge Base', level=3)
doc.add_paragraph('The platform leverages Google Document AI for processing research documents, news, and institutional reports:')
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Component'
hdr[1].text = 'Service'
hdr[2].text = 'Purpose'
data = [
    ('PDF Processor', 'Document AI', 'Extract text from research PDFs'),
    ('Embedding Generator', 'Vertex AI Embeddings', 'Create vector representations'),
    ('Vector Store', 'Vertex AI Vector Search', 'Semantic search over documents'),
    ('Knowledge Graph', 'Enterprise Knowledge Graph', 'Entity relationships'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

doc.add_paragraph()

# Part 3: Implementation Phases
doc.add_heading('Part 3: Implementation Phases', level=1)

doc.add_heading('3.1 Phase 1: Trader-Focused AI Alert Engine', level=2)
doc.add_paragraph('Focus: Real-time alerts for equities and cryptocurrency')

doc.add_heading('MVP Symbol Universe', level=3)
table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Category'
hdr[1].text = 'Examples'
hdr[2].text = 'Count'
data = [
    ('US Equities', 'AAPL, NVDA, JPM, LLY, MSFT, GOOGL, AMZN, META, TSLA, etc.', '50'),
    ('ETFs', 'SPY, QQQ, IWM, DIA, VTI, VOO, GLD, etc.', '30'),
    ('Cryptocurrency', 'BTC/USD, ETH/USD, SOL/USD, XRP/USD, etc.', '20'),
    ('Forex', 'EUR/USD, GBP/USD, USD/JPY, etc.', '20'),
    ('Indices', 'SPX, NDX, DJI, VIX, etc.', '14'),
    ('Commodities', 'XAU/USD, XAG/USD, CL, NG, etc.', '16'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

doc.add_paragraph()

doc.add_heading('Phase 1 Deliverables', level=3)
doc.add_paragraph('1. Real-Time Alert System')
doc.add_paragraph('   - Trend flip alerts')
doc.add_paragraph('   - Momentum reversal signals')
doc.add_paragraph('   - Volatility expansion warnings')
doc.add_paragraph('   - Breakout/breakdown detection')
doc.add_paragraph('   - Pattern confirmation alerts')

doc.add_paragraph('2. AI-Enhanced Dashboards')
doc.add_paragraph('   - Regime classification display')
doc.add_paragraph('   - Volatility heatmaps')
doc.add_paragraph('   - Correlation monitors')
doc.add_paragraph('   - Sentiment indicators')

doc.add_paragraph('3. LLM Commentary (Basic)')
doc.add_paragraph('   - Alert explanation generation')
doc.add_paragraph('   - Daily market summaries')
doc.add_paragraph('   - Pattern interpretation')

doc.add_heading('3.2 Phase 2: Institutional Treasury Intelligence System', level=2)
doc.add_paragraph('Focus: Enterprise-grade macro and treasury intelligence')
table = doc.add_table(rows=7, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Capability'
hdr[1].text = 'Description'
data = [
    ('Cross-Asset Correlation Engine', 'Real-time correlation matrices across all asset classes'),
    ('Macro Regime Classification', 'Risk-on/off detection, volatility regime scoring'),
    ('Treasury Exposure Analytics', 'FX risk mapping, yield sensitivity analysis'),
    ('Yield Curve Analytics', 'Curve shape analysis, rate sensitivity'),
    ('Early Warning System', 'Liquidity stress, correlation breakdown alerts'),
    ('AI-Generated Institutional Reports', 'Compliance-ready research commentary'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]

doc.add_paragraph()

# Part 4: GCP Infrastructure
doc.add_heading('Part 4: GCP Infrastructure', level=1)

doc.add_heading('4.1 Service Architecture', level=2)
doc.add_paragraph('GCP PROJECT: cryptobot-462709')
doc.add_paragraph('')
doc.add_paragraph('The infrastructure consists of:')
doc.add_paragraph('- Cloud Run (Frontend + API)')
doc.add_paragraph('- BigQuery (Data Warehouse)')
doc.add_paragraph('- Vertex AI (ML + Gemini 3)')
doc.add_paragraph('- Cloud Functions (Data Fetch)')
doc.add_paragraph('- Cloud Storage (Documents)')
doc.add_paragraph('- Cloud Scheduler (Triggers)')

doc.add_heading('4.2 Cost Optimization Strategy', level=2)
table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Service'
hdr[1].text = 'Optimization'
hdr[2].text = 'Monthly Estimate'
data = [
    ('BigQuery', 'Partitioned tables, clustering', '$50-100'),
    ('Cloud Functions', 'Efficient scheduling, cold start optimization', '$130'),
    ('Cloud Run', 'Auto-scaling, min instances', '$50-100'),
    ('Vertex AI', 'Batch predictions, model caching', '$200-500'),
    ('Gemini 3', 'Caching, prompt optimization', '$100-300'),
    ('Total Estimate', '', '$530-1,130/month'),
]
for i, row_data in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

doc.add_paragraph()

# Part 5: Missing Fields Implementation Plan
doc.add_heading('Part 5: Missing Fields Implementation Plan', level=1)
doc.add_paragraph('Based on Phase 1 Methodology analysis, the following fields need to be added:')

doc.add_heading('Priority 1: Critical ML Features (Add Immediately)', level=2)
doc.add_paragraph('Log returns: weekly_log_return - ln(close_t / close_t-1)')
doc.add_paragraph('Multi-lag returns: return_2w, return_4w')
doc.add_paragraph('RSI enhancements: rsi_slope, rsi_zscore, rsi_overbought_flag, rsi_oversold_flag')
doc.add_paragraph('MACD enhancements: macd_hist, macd_cross_flag')
doc.add_paragraph('EMA suite: ema_5, ema_10, ema_20, ema_50, ema_100, ema_200')
doc.add_paragraph('MA distance features: close_vs_sma20_pct, close_vs_sma50_pct, close_vs_sma200_pct, close_vs_ema20_pct')
doc.add_paragraph('EMA slopes: ema_slope_20, ema_slope_50')

doc.add_heading('Priority 2: Enhanced Features (Add Next Sprint)', level=2)
doc.add_paragraph('ATR enhancements: atr_pct, atr_zscore, atr_slope')
doc.add_paragraph('Volume enhancements: volume_zscore, volume_ratio')
doc.add_paragraph('ADX trend: adx, plus_di, minus_di')
doc.add_paragraph('Candle geometry: candle_body_pct, candle_range_pct')

doc.add_heading('Priority 3: Advanced Structure (Phase 2)', level=2)
doc.add_paragraph('Pivot structure: pivot_high_flag, pivot_low_flag, dist_to_pivot_high, dist_to_pivot_low')
doc.add_paragraph('Regime classification: regime_state, regime_confidence')

doc.add_paragraph()

# Part 6: Implementation Roadmap
doc.add_heading('Part 6: Implementation Roadmap', level=1)

doc.add_heading('Week 1-2: Data Schema Enhancement', level=2)
doc.add_paragraph('- Add Priority 1 missing fields to BigQuery tables')
doc.add_paragraph('- Update data fetching scripts with new calculations')
doc.add_paragraph('- Backfill historical data with new indicators')

doc.add_heading('Week 3-4: ML Pipeline Setup', level=2)
doc.add_paragraph('- Configure Vertex AI Feature Store')
doc.add_paragraph('- Create initial training pipelines')
doc.add_paragraph('- Deploy baseline models')

doc.add_heading('Week 5-6: LLM Integration', level=2)
doc.add_paragraph('- Implement RAG pipeline with BigQuery')
doc.add_paragraph('- Configure Gemini 3 Pro prompts')
doc.add_paragraph('- Build commentary generation endpoints')

doc.add_heading('Week 7-8: Testing & Optimization', level=2)
doc.add_paragraph('- Backtest ML models')
doc.add_paragraph('- Evaluate LLM output quality')
doc.add_paragraph('- Performance optimization')

doc.add_paragraph()

# Appendix
doc.add_heading('Appendix A: Technical Indicator Formulas', level=1)

doc.add_heading('A.1 Log Return', level=2)
doc.add_paragraph('weekly_log_return = ln(close_t) - ln(close_t-1)')

doc.add_heading('A.2 RSI Enhancements', level=2)
doc.add_paragraph('rsi_slope = rsi_t - rsi_t-1')
doc.add_paragraph('rsi_zscore = (rsi - mean_100w) / std_100w')
doc.add_paragraph('rsi_overbought_flag = 1 if rsi > 70 else 0')
doc.add_paragraph('rsi_oversold_flag = 1 if rsi < 30 else 0')

doc.add_heading('A.3 MACD Cross Flag', level=2)
doc.add_paragraph('macd_cross_flag = +1 if macd crosses above signal')
doc.add_paragraph('                = -1 if macd crosses below signal')
doc.add_paragraph('                = 0 otherwise')

doc.add_heading('A.4 MA Distance', level=2)
doc.add_paragraph('close_vs_sma20_pct = (close / sma_20 - 1) * 100')

doc.add_heading('A.5 EMA Slope', level=2)
doc.add_paragraph('ema_slope_20 = ema_20_t - ema_20_t-1')

doc.add_heading('A.6 ATR Percentage', level=2)
doc.add_paragraph('atr_pct = (atr_14 / close) * 100')

doc.add_heading('A.7 Volume Z-Score', level=2)
doc.add_paragraph('volume_zscore = (volume - mean_20) / std_20')

doc.add_paragraph()

# Document Control
doc.add_heading('Document Control', level=1)
table = doc.add_table(rows=3, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Version'
hdr[1].text = 'Date'
hdr[2].text = 'Author'
hdr[3].text = 'Changes'
row1 = table.rows[1].cells
row1[0].text = '1.0'
row1[1].text = 'Nov 2025'
row1[2].text = 'System'
row1[3].text = 'Initial draft'
row2 = table.rows[2].cells
row2[0].text = '2.0'
row2[1].text = 'Nov 2025'
row2[2].text = 'Claude/AI'
row2[3].text = 'Combined Phase 1 + AI/LLM specs'

doc.add_paragraph()
doc.add_paragraph('End of Document')

# Save the document
doc.save('FINTECH_AI_DATA_ARCHITECTURE_MASTER_SPECIFICATION.docx')
print('Document created successfully!')
