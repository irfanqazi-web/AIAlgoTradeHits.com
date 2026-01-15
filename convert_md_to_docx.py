"""
Convert all Markdown documentation files to Word .docx format
"""
import os
import sys
import io
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Fix Windows encoding
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Style for hyperlink
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._element.append(hyperlink)
    return hyperlink

def parse_markdown_to_docx(md_file, docx_file):
    """Convert a markdown file to a Word document"""
    print(f"Converting: {os.path.basename(md_file)}")

    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_language = ''
    code_lines = []
    in_table = False
    table = None

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Code blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text, style='Normal')
                p_format = p.paragraph_format
                p_format.left_indent = Inches(0.5)
                p_format.space_before = Pt(6)
                p_format.space_after = Pt(6)

                for run in p.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 0)

                # Light gray background
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'F5F5F5')
                p._element.get_or_add_pPr().append(shading_elm)

                in_code_block = False
                code_lines = []
                code_language = ''
            else:
                # Start of code block
                in_code_block = True
                code_language = line[3:].strip()
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Empty lines
        if not line.strip():
            if not in_table:
                doc.add_paragraph()
            i += 1
            continue

        # Headers
        if line.startswith('#'):
            in_table = False
            header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if header_match:
                level = len(header_match.group(1))
                text = header_match.group(2).strip()

                # Remove markdown formatting from header text
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # Bold
                text = re.sub(r'\*(.+?)\*', r'\1', text)  # Italic
                text = re.sub(r'`(.+?)`', r'\1', text)  # Code

                p = doc.add_paragraph(text)
                if level == 1:
                    p.style = 'Heading 1'
                elif level == 2:
                    p.style = 'Heading 2'
                elif level == 3:
                    p.style = 'Heading 3'
                elif level == 4:
                    p.style = 'Heading 4'
                else:
                    p.style = 'Heading 5'
            i += 1
            continue

        # Horizontal rules
        if line.strip() in ['---', '***', '___']:
            in_table = False
            p = doc.add_paragraph('_' * 80)
            run = p.runs[0]
            run.font.color.rgb = RGBColor(200, 200, 200)
            i += 1
            continue

        # Tables
        if '|' in line and not in_table:
            # Start of table
            in_table = True
            headers = [cell.strip() for cell in line.split('|')[1:-1]]

            # Skip separator line
            i += 1
            if i < len(lines) and re.match(r'^\|[\s\-:|]+\|$', lines[i]):
                i += 1

            # Count rows
            row_count = 0
            j = i
            while j < len(lines) and '|' in lines[j] and not lines[j].startswith('#'):
                row_count += 1
                j += 1

            # Create table
            table = doc.add_table(rows=row_count + 1, cols=len(headers))
            table.style = 'Light Grid Accent 1'

            # Add headers
            for col_idx, header in enumerate(headers):
                cell = table.rows[0].cells[col_idx]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            # Add rows
            row_idx = 1
            while i < len(lines) and '|' in lines[i] and not lines[i].startswith('#'):
                cells = [cell.strip() for cell in lines[i].split('|')[1:-1]]
                for col_idx, cell_text in enumerate(cells):
                    if col_idx < len(headers):
                        # Remove markdown formatting
                        cell_text = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
                        cell_text = re.sub(r'\*(.+?)\*', r'\1', cell_text)
                        cell_text = re.sub(r'`(.+?)`', r'\1', cell_text)

                        table.rows[row_idx].cells[col_idx].text = cell_text
                row_idx += 1
                i += 1

            in_table = False
            continue

        # Bullet lists
        if line.startswith('- ') or line.startswith('* ') or re.match(r'^\d+\.\s', line):
            in_table = False
            text = re.sub(r'^[-*]\s+', '', line)
            text = re.sub(r'^\d+\.\s+', '', text)

            # Parse inline formatting
            p = doc.add_paragraph(style='List Bullet' if line[0] in '-*' else 'List Number')

            # Process inline formatting
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))', text)
            for part in parts:
                if not part:
                    continue

                # Bold
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                # Italic
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.font.italic = True
                # Code
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                # Links
                elif part.startswith('[') and '](' in part:
                    match = re.match(r'\[(.+?)\]\((.+?)\)', part)
                    if match:
                        link_text = match.group(1)
                        url = match.group(2)
                        add_hyperlink(p, link_text, url)
                else:
                    p.add_run(part)

            i += 1
            continue

        # Regular paragraphs
        in_table = False
        p = doc.add_paragraph()

        # Process inline formatting
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))', line)
        for part in parts:
            if not part:
                continue

            # Bold
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.font.bold = True
            # Italic
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                run = p.add_run(part[1:-1])
                run.font.italic = True
            # Code
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(200, 0, 0)
            # Links
            elif part.startswith('[') and '](' in part:
                match = re.match(r'\[(.+?)\]\((.+?)\)', part)
                if match:
                    link_text = match.group(1)
                    url = match.group(2)
                    add_hyperlink(p, link_text, url)
            else:
                p.add_run(part)

        i += 1

    # Save document
    doc.save(docx_file)
    print(f"  ✓ Created: {os.path.basename(docx_file)}")

def main():
    """Convert all markdown files in current directory to .docx"""
    current_dir = Path(__file__).parent

    # Find all .md files
    md_files = list(current_dir.glob('*.md'))

    print(f"\nFound {len(md_files)} markdown files\n")
    print("=" * 80)

    converted = 0
    skipped = 0
    failed = 0

    for md_file in md_files:
        try:
            # Skip README.md
            if md_file.name == 'README.md':
                print(f"Skipping: {md_file.name} (README)")
                skipped += 1
                continue

            docx_file = md_file.with_suffix('.docx')

            # Skip if .docx already exists
            if docx_file.exists():
                print(f"Skipping: {md_file.name} (already has .docx)")
                skipped += 1
                continue

            parse_markdown_to_docx(str(md_file), str(docx_file))
            converted += 1
        except Exception as e:
            print(f"  ✗ ERROR: {md_file.name}")
            print(f"    {str(e)}")
            failed += 1

    print("=" * 80)
    print(f"\nConversion complete!")
    print(f"✓ Converted: {converted}")
    print(f"⊘ Skipped (already exists): {skipped}")
    print(f"✗ Failed: {failed}")
    print(f"Total files: {len(md_files)}")

if __name__ == '__main__':
    main()
