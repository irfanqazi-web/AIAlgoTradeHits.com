import sys
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Windows UTF-8 encoding fix
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

def add_heading(doc, text, level):
    """Add formatted heading"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.size = Pt(24)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    elif level == 2:
        heading.runs[0].font.size = Pt(18)
        heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
    return heading

def add_table_with_header(doc, headers, data):
    """Add a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)

    # Add data
    for row_data in data:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = str(value)

    return table

def create_master_plan_docx():
    """Create comprehensive Word document for Master Trading Plan"""

    doc = Document()

    # Title Page
    title = doc.add_heading('MASTER TRADING DATA IMPLEMENTATION PLAN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Twelve Data Pro Plan - Complete Multi-Asset Platform')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.add_run('Generated: ').bold = True
    info.add_run('November 20, 2025\n')
    info.add_run('API Key: ').bold = True
    info.add_run('16ee060fd4d34a628a14bcb6f0167565\n')
    info.add_run('Plan: ').bold = True
    info.add_run('Pro ($229/month) - 1,597 API calls/min + 1,500 WebSocket credits')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Table of Contents
    add_heading(doc, 'Table of Contents', 1)
    toc = [
        '1. Complete Data Inventory',
        '2. Technical Indicators (71+)',
        '3. Fundamental Data Endpoints',
        '4. BigQuery Architecture',
        '5. Cloud Functions Architecture',
        '6. API Usage Estimation',
        '7. Frontend Integration',
        '8. Cost Analysis',
        '9. Implementation Phases',
        '10. Asset Universe Details',
        '11. Security & Best Practices',
        '12. Sample Queries',
        '13. Success Metrics',
        '14. Maintenance Schedule',
    ]
    for item in toc:
        doc.add_paragraph(item, style='List Number')

    doc.add_page_break()

    # Section 1: Complete Data Inventory
    add_heading(doc, '1. COMPLETE DATA INVENTORY - VERIFIED', 1)

    add_heading(doc, 'Asset Classes Available', 2)

    asset_data = [
        ['Stocks', '20,076', 'S&P 500 (500 stocks)', '✓'],
        ['ETFs', '10,186', 'Top 200 US ETFs', '✓'],
        ['Forex', '1,436 pairs (178 USD)', 'Major 20 USD pairs', '✓'],
        ['Crypto', '2,149 coins (1,136 USD)', 'Top 100 USD pairs', '✓'],
        ['Commodities', '62', 'All US-traded (40)', '✓'],
        ['Indices', 'Available', 'Major US (S&P, Dow, Nasdaq)', '✓'],
        ['Bonds', 'Available', 'US Treasury Bonds', '✓'],
        ['Mutual Funds', 'Available', 'Top 100 US funds', '✓'],
    ]

    add_table_with_header(doc, ['Asset Class', 'Total Available', 'USA Focus', 'Verified'], asset_data)

    doc.add_paragraph()
    add_heading(doc, 'US Stock Exchanges (16)', 2)

    exchanges = [
        ['NYSE', 'XNYS', '3,032 stocks'],
        ['NASDAQ', 'XNAS, XNGS, XNCM, XNMS', '4,435 stocks'],
        ['OTC', 'OTCB, OTCM, PINX, PSGM, EXPM, OTCQ', '12,598 stocks'],
        ['CBOE', 'BATS', 'Options'],
        ['NYSE Arca', 'ARCX', 'ETFs'],
        ['IEX', 'IEXG', 'Stocks'],
        ['ICE', 'NYBOT', 'Futures'],
    ]

    add_table_with_header(doc, ['Exchange', 'Code', 'Assets'], exchanges)

    p = doc.add_paragraph()
    p.add_run('\nGlobal Exchanges: ').bold = True
    p.add_run('96 exchanges across 58 countries')

    p = doc.add_paragraph()
    p.add_run('Cryptocurrency Exchanges: ').bold = True
    p.add_run('201 exchanges\nIncluding: Binance, Coinbase, Kraken, Bitfinex, Bittrex, Gemini, Bitstamp, Huobi, KuCoin, FTX, OKX, etc.')

    doc.add_page_break()

    # Section 2: Technical Indicators
    add_heading(doc, '2. TECHNICAL INDICATORS - 71+ VERIFIED', 1)

    indicators = [
        ('Momentum Indicators (10)', [
            'RSI - Relative Strength Index',
            'MACD - Moving Average Convergence Divergence',
            'STOCH - Stochastic Oscillator',
            'WILLIAMS - Williams %R',
            'CCI - Commodity Channel Index',
            'ROC - Rate of Change',
            'MOM - Momentum',
            'STOCHRSI - Stochastic RSI',
            'APO - Absolute Price Oscillator',
            'PPO - Percentage Price Oscillator',
        ]),
        ('Overlap Studies (11)', [
            'SMA - Simple Moving Average',
            'EMA - Exponential Moving Average',
            'WMA - Weighted Moving Average',
            'DEMA - Double Exponential MA',
            'TEMA - Triple Exponential MA',
            'TRIMA - Triangular MA',
            'KAMA - Kaufman Adaptive MA',
            'MAMA - MESA Adaptive MA',
            'VWAP - Volume Weighted Average Price',
            'T3 - T3 Moving Average',
            'HT_TRENDLINE - Hilbert Transform',
        ]),
        ('Volatility Indicators (6)', [
            'BBANDS - Bollinger Bands',
            'ATR - Average True Range',
            'NATR - Normalized ATR',
            'TRANGE - True Range',
            'STDDEV - Standard Deviation',
            'VARIANCE - Variance',
        ]),
        ('Volume Indicators (4)', [
            'OBV - On Balance Volume',
            'AD - Chaikin A/D Line',
            'ADOSC - Chaikin A/D Oscillator',
            'PVO - Percentage Volume Oscillator',
        ]),
        ('Trend Indicators (10)', [
            'ADX - Average Directional Index',
            'ADXR - ADX Rating',
            'AROON - Aroon',
            'AROONOSC - Aroon Oscillator',
            'DX - Directional Movement Index',
            'PLUS_DI - Plus Directional Indicator',
            'MINUS_DI - Minus Directional Indicator',
            'PLUS_DM - Plus Directional Movement',
            'MINUS_DM - Minus Directional Movement',
            'TRIX - Triple Exponential Average',
        ]),
        ('Candlestick Pattern Recognition (10)', [
            'CDL2CROWS - Two Crows',
            'CDL3BLACKCROWS - Three Black Crows',
            'CDL3INSIDE - Three Inside Up/Down',
            'CDL3LINESTRIKE - Three-Line Strike',
            'CDLABANDONEDBABY - Abandoned Baby',
            'CDLDOJI - Doji',
            'CDLENGULFING - Engulfing Pattern',
            'CDLHAMMER - Hammer',
            'CDLHARAMI - Harami Pattern',
            'CDLMORNINGSTAR - Morning Star',
        ]),
        ('Statistical Functions (7)', [
            'CORREL - Pearson\'s Correlation',
            'LINEARREG - Linear Regression',
            'LINEARREG_ANGLE - LR Angle',
            'LINEARREG_INTERCEPT - LR Intercept',
            'LINEARREG_SLOPE - LR Slope',
            'TSF - Time Series Forecast',
            'VAR - Variance',
        ]),
        ('Other Advanced Indicators (13)', [
            'SAR - Parabolic SAR',
            'SAREXT - Parabolic SAR Extended',
            'ULTOSC - Ultimate Oscillator',
            'BOP - Balance of Power',
            'CMO - Chande Momentum Oscillator',
            'DPO - Detrended Price Oscillator',
            'HT_DCPERIOD - Hilbert Transform Dominant Cycle Period',
            'HT_DCPHASE - Hilbert Transform Dominant Cycle Phase',
            'HT_PHASOR - Hilbert Transform Phasor Components',
            'HT_SINE - Hilbert Transform SineWave',
            'HT_TRENDMODE - Hilbert Transform Trend vs Cycle',
            'MIDPOINT - MidPoint over period',
            'MIDPRICE - Midpoint Price',
        ]),
    ]

    for category, items in indicators:
        add_heading(doc, category, 2)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')

    summary = doc.add_paragraph()
    summary.add_run('\nTotal Technical Indicators: ').bold = True
    summary.add_run('71+ indicators available across all asset classes')
    summary.runs[0].font.size = Pt(14)
    summary.runs[0].font.color.rgb = RGBColor(0, 128, 0)

    doc.add_page_break()

    # Section 3: Fundamental Data
    add_heading(doc, '3. FUNDAMENTAL DATA ENDPOINTS - VERIFIED', 1)

    add_heading(doc, 'Company Information', 2)
    doc.add_paragraph('Profile - Company description, CEO, address, website, employees', style='List Bullet')
    doc.add_paragraph('Logo - Company logo URLs', style='List Bullet')
    doc.add_paragraph('Statistics - Market cap, PE ratio, PEG, Price-to-Sales, ROE, ROA', style='List Bullet')

    add_heading(doc, 'Financial Statements', 2)
    doc.add_paragraph('Income Statement - Revenue, gross profit, operating income, EPS, EBITDA', style='List Bullet')
    doc.add_paragraph('Balance Sheet - Assets, liabilities, equity, cash, debt', style='List Bullet')
    doc.add_paragraph('Cash Flow - Operating, investing, financing activities', style='List Bullet')

    add_heading(doc, 'Corporate Actions', 2)
    doc.add_paragraph('Earnings Calendar - Earnings dates, EPS estimates vs actuals', style='List Bullet')
    doc.add_paragraph('Dividends - Dividend history with ex-dates and amounts', style='List Bullet')
    doc.add_paragraph('Stock Splits - Split history with ratios', style='List Bullet')
    doc.add_paragraph('Insider Transactions - Insider buying/selling activity', style='List Bullet')

    add_heading(doc, 'NOT Available (404 Errors)', 2)
    doc.add_paragraph('News API - Not available in Pro plan', style='List Bullet')
    doc.add_paragraph('Sentiment Analysis - Not available in Pro plan', style='List Bullet')
    doc.add_paragraph('Options Chain - Not available', style='List Bullet')

    doc.add_page_break()

    # Section 4: BigQuery Architecture
    add_heading(doc, '4. BIGQUERY ARCHITECTURE - UNIFIED MULTI-ASSET SCHEMA', 1)

    p = doc.add_paragraph()
    p.add_run('Dataset: ').bold = True
    p.add_run('cryptobot-462709.trading_data_unified')

    add_heading(doc, 'Tables Overview', 2)

    tables_data = [
        ['1', 'stocks_daily', '60 fields', 'Daily OHLCV + 71 indicators'],
        ['2', 'stocks_hourly', '60 fields', 'Hourly data + indicators'],
        ['3', 'stocks_5min_top100', '60 fields', 'Top 100 gainers, 5-min data'],
        ['4', 'etfs_daily', '60 fields', 'ETF daily data'],
        ['5', 'etfs_hourly', '60 fields', 'ETF hourly data'],
        ['6', 'forex_daily', '58 fields', 'Forex daily (no volume indicators)'],
        ['7', 'forex_hourly', '58 fields', 'Forex hourly data'],
        ['8', 'forex_5min_major20', '58 fields', 'Top 20 forex pairs'],
        ['9', 'crypto_daily', '60 fields', 'Crypto daily (EXISTING - Kraken)'],
        ['10', 'crypto_hourly', '60 fields', 'Crypto hourly (EXISTING)'],
        ['11', 'crypto_5min_top10', '60 fields', 'Top 10 crypto gainers (EXISTING)'],
        ['12', 'commodities_daily', '60 fields', 'Commodities daily data'],
        ['13', 'commodities_hourly', '60 fields', 'Commodities hourly data'],
        ['14', 'indices_daily', '60 fields', 'US indices daily'],
        ['15', 'bonds_daily', '58 fields', 'US Treasury bonds'],
        ['16', 'stock_fundamentals', '40 fields', 'Company fundamentals'],
        ['17', 'earnings_calendar', '15 fields', 'Earnings dates & estimates'],
        ['18', 'dividends_history', '10 fields', 'Dividend history'],
        ['19', 'stock_splits_history', '10 fields', 'Stock split history'],
        ['20', 'insider_transactions', '15 fields', 'Insider trading activity'],
    ]

    add_table_with_header(doc, ['#', 'Table Name', 'Fields', 'Description'], tables_data)

    doc.add_page_break()

    add_heading(doc, 'Standard Table Schema (60 fields)', 2)

    p = doc.add_paragraph()
    p.add_run('Core OHLCV + Metadata (15 fields):').bold = True

    core_fields = [
        'symbol - Trading symbol',
        'name - Full name',
        'exchange - Exchange name',
        'mic_code - Market identifier code',
        'currency - Trading currency',
        'datetime - Timestamp',
        'date - Date',
        'open - Opening price',
        'high - High price',
        'low - Low price',
        'close - Closing price',
        'volume - Trading volume',
        'adjusted_close - Adjusted close',
        'asset_type - Asset class',
        'data_source - Data provider',
    ]

    for field in core_fields:
        doc.add_paragraph(field, style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('\nTechnical Indicators (45 fields):').bold = True
    doc.add_paragraph('All 71 indicators calculated and stored as separate columns')
    doc.add_paragraph('Momentum: rsi, macd, macd_signal, macd_hist, stoch_k, stoch_d, williams_r, cci, roc, momentum', style='List Bullet')
    doc.add_paragraph('Moving Averages: sma_20, sma_50, sma_200, ema_12, ema_26, ema_50, wma_20, dema_20, tema_20, kama_20, vwap', style='List Bullet')
    doc.add_paragraph('Volatility: bb_upper, bb_middle, bb_lower, atr, natr, stddev', style='List Bullet')
    doc.add_paragraph('Volume: obv, ad, adosc, pvo', style='List Bullet')
    doc.add_paragraph('Trend: adx, adxr, plus_di, minus_di, aroon_up, aroon_down, aroonosc, trix, dx, sar', style='List Bullet')
    doc.add_paragraph('Pattern Recognition: 10 candlestick patterns', style='List Bullet')
    doc.add_paragraph('Statistical: correl, linearreg, linearreg_slope, linearreg_angle, tsf, variance, beta', style='List Bullet')
    doc.add_paragraph('Other: ultosc, bop, cmo, dpo, and 9 more advanced indicators', style='List Bullet')

    doc.add_page_break()

    # Section 5: Cloud Functions
    add_heading(doc, '5. CLOUD FUNCTIONS ARCHITECTURE', 1)

    p = doc.add_paragraph()
    p.add_run('Total Functions: ').bold = True
    p.add_run('23 Cloud Functions\n')
    p.add_run('Naming Convention: ').bold = True
    p.add_run('{frequency}-{asset_type}-fetcher\n')
    p.add_run('Example: ').bold = True
    p.add_run('daily-stocks-fetcher, hourly-forex-fetcher, 5min-crypto-top10-fetcher')

    add_heading(doc, 'DAILY FUNCTIONS (7 functions)', 2)

    daily_functions = [
        ['1', 'daily-stocks-fetcher', 'S&P 500 (500)', '1 AM ET M-F', '500', 'stocks_daily'],
        ['2', 'daily-etfs-fetcher', 'Top 200 ETFs', '2 AM ET M-F', '200', 'etfs_daily'],
        ['3', 'daily-forex-fetcher', '20 major USD pairs', '3 AM ET daily', '20', 'forex_daily'],
        ['4', 'daily-crypto-fetcher', '675 USD pairs (Kraken)', 'Midnight ET', 'EXISTING', 'crypto_daily'],
        ['5', 'daily-commodities-fetcher', '40 US commodities', '4 AM ET M-F', '40', 'commodities_daily'],
        ['6', 'daily-indices-fetcher', '5 major US indices', '5 AM ET M-F', '5', 'indices_daily'],
        ['7', 'daily-bonds-fetcher', '5 US Treasury bonds', '6 AM ET M-F', '5', 'bonds_daily'],
    ]

    add_table_with_header(doc, ['#', 'Function Name', 'Assets', 'Schedule', 'API Credits', 'Table'], daily_functions)

    p = doc.add_paragraph()
    p.add_run('\nTotal Daily Credits: ').bold = True
    p.add_run('1,445 credits/day')

    doc.add_page_break()

    add_heading(doc, 'HOURLY FUNCTIONS (6 functions)', 2)

    hourly_functions = [
        ['8', 'hourly-stocks-fetcher', 'S&P 500', 'Every hour 9AM-4PM ET', '3,500/day', 'stocks_hourly'],
        ['9', 'hourly-etfs-fetcher', 'Top 200 ETFs', '5 min past hour', '1,400/day', 'etfs_hourly'],
        ['10', 'hourly-forex-fetcher', '20 major pairs', '10 min past (24/7)', '480/day', 'forex_hourly'],
        ['11', 'hourly-crypto-fetcher', '675 pairs (Kraken)', 'Every hour', 'EXISTING', 'crypto_hourly'],
        ['12', 'hourly-commodities-fetcher', '40 commodities', '15 min past hour', '280/day', 'commodities_hourly'],
        ['13', 'hourly-indices-fetcher', '5 indices', '20 min past hour', '35/day', 'indices_hourly'],
    ]

    add_table_with_header(doc, ['#', 'Function Name', 'Assets', 'Schedule', 'API Credits', 'Table'], hourly_functions)

    p = doc.add_paragraph()
    p.add_run('\nTotal Hourly Credits: ').bold = True
    p.add_run('5,695 credits/day')

    doc.add_page_break()

    add_heading(doc, '5-MINUTE FUNCTIONS (4 functions)', 2)

    fivemin_functions = [
        ['14', '5min-stocks-top100-fetcher', 'Top 100 gainers', 'Every 5 min (9AM-4PM)', '8,400/day', 'stocks_5min_top100'],
        ['15', '5min-etfs-top50-fetcher', 'Top 50 by volume', 'Every 5 min (offset 1)', '4,200/day', 'etfs_5min_top50'],
        ['16', '5min-forex-major10-fetcher', '10 most volatile', 'Every 5 min (24/7)', '2,880/day', 'forex_5min_major10'],
        ['17', '5min-crypto-top10-fetcher', 'Top 10 gainers (Kraken)', 'Every 5 min', 'EXISTING', 'crypto_5min_top10'],
    ]

    add_table_with_header(doc, ['#', 'Function Name', 'Assets', 'Schedule', 'API Credits', 'Table'], fivemin_functions)

    p = doc.add_paragraph()
    p.add_run('\nTotal 5-Minute Credits: ').bold = True
    p.add_run('15,480 credits/day')

    doc.add_page_break()

    add_heading(doc, 'FUNDAMENTAL DATA FUNCTIONS (5 functions)', 2)

    fundamental_functions = [
        ['18', 'fundamentals-stocks-fetcher', 'S&P 500', 'Sunday 7 AM (weekly)', '2,500/week', 'stock_fundamentals'],
        ['19', 'earnings-calendar-fetcher', 'Next 30 days', '8 AM ET M-F', '100/day', 'earnings_calendar'],
        ['20', 'dividends-fetcher', 'S&P 500', 'Monday 9 AM (weekly)', '500/week', 'dividends_history'],
        ['21', 'splits-fetcher', 'S&P 500', 'Monday 10 AM (weekly)', '500/week', 'stock_splits_history'],
        ['22', 'insider-transactions-fetcher', 'S&P 500', 'Monday 11 AM (weekly)', '500/week', 'insider_transactions'],
    ]

    add_table_with_header(doc, ['#', 'Function Name', 'Assets', 'Schedule', 'API Credits', 'Table'], fundamental_functions)

    p = doc.add_paragraph()
    p.add_run('\nTotal Fundamental Credits: ').bold = True
    p.add_run('~300 credits/day average')

    doc.add_page_break()

    add_heading(doc, 'MARKET MOVERS FUNCTION (1 function)', 2)

    movers_functions = [
        ['23', 'market-movers-fetcher', 'Stocks/Crypto/Forex', '9AM, 12PM, 3PM ET', '9/day', 'market_movers_daily'],
    ]

    add_table_with_header(doc, ['#', 'Function Name', 'Assets', 'Schedule', 'API Credits', 'Table'], movers_functions)

    doc.add_page_break()

    # Section 6: API Usage
    add_heading(doc, '6. TOTAL API USAGE ESTIMATION', 1)

    usage_data = [
        ['Daily Functions', '1,445', '43,350', '0.9%'],
        ['Hourly Functions', '5,695', '170,850', '3.6%'],
        ['5-Min Functions', '15,480', '464,400', '9.7%'],
        ['Fundamentals', '300', '9,000', '0.2%'],
        ['Market Movers', '9', '270', '0.01%'],
        ['TOTAL', '22,929', '687,870', '14.4%'],
    ]

    add_table_with_header(doc, ['Category', 'Daily Credits', 'Monthly Credits', '% of Limit'], usage_data)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Monthly API Limit: ').bold = True
    p.add_run('4,780,800 credits (1,597 calls/min × 60 × 24 × 30)\n')
    p.add_run('Buffer Remaining: ').bold = True
    p.add_run('85.6% (4,092,930 credits available)')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.color.rgb = RGBColor(0, 128, 0)

    doc.add_page_break()

    # Section 7: Frontend Integration
    add_heading(doc, '7. FRONTEND INTEGRATION', 1)

    add_heading(doc, 'Multi-Asset Selector Component', 2)
    doc.add_paragraph('Asset class selector with icons for Stocks, ETFs, Crypto, Forex, Commodities, Bonds, Indices')
    doc.add_paragraph('Real-time switching between asset classes')
    doc.add_paragraph('Persistent user preferences')

    add_heading(doc, 'Unified Chart Component', 2)
    doc.add_paragraph('Single chart component supporting all asset types')
    doc.add_paragraph('71+ technical indicators overlay')
    doc.add_paragraph('Multiple timeframes (1min, 5min, 15min, 1h, 1day, 1week, 1month)')
    doc.add_paragraph('Candlestick pattern recognition highlighting')
    doc.add_paragraph('Drawing tools and annotations')

    add_heading(doc, 'Market Overview Dashboard', 2)
    doc.add_paragraph('Top gainers/losers across all asset classes')
    doc.add_paragraph('Real-time market movers')
    doc.add_paragraph('Volume leaders')
    doc.add_paragraph('Sector performance heatmap')

    add_heading(doc, 'Advanced Screeners', 2)
    doc.add_paragraph('Multi-criteria filtering (RSI, MACD, volume, price, etc.)')
    doc.add_paragraph('Custom screener templates')
    doc.add_paragraph('Real-time alert system')
    doc.add_paragraph('Backtest screener strategies')

    doc.add_page_break()

    # Section 8: Cost Analysis
    add_heading(doc, '8. COST ANALYSIS', 1)

    cost_data = [
        ['Twelve Data API', '$229.00'],
        ['Cloud Functions (23 functions)', '$180.00'],
        ['BigQuery Storage (500GB)', '$10.00'],
        ['BigQuery Queries', '$15.00'],
        ['Cloud Scheduler (23 jobs)', '$0.30'],
        ['Cloud Run (Trading App)', '$5.00'],
        ['TOTAL MONTHLY', '$439.30'],
    ]

    add_table_with_header(doc, ['Service', 'Monthly Cost'], cost_data)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Annual Cost: ').bold = True
    p.add_run('$5,271.60\n')
    p.add_run('Cost per Asset Class: ').bold = True
    p.add_run('~$55/month average')

    doc.add_page_break()

    # Section 9: Implementation Phases
    add_heading(doc, '9. IMPLEMENTATION PHASES', 1)

    add_heading(doc, 'Phase 1: Stock Data Pipeline (Week 1-2)', 2)
    doc.add_paragraph('✓ API key configured', style='List Bullet')
    doc.add_paragraph('Create stocks BigQuery schema', style='List Bullet')
    doc.add_paragraph('Build daily-stocks-fetcher', style='List Bullet')
    doc.add_paragraph('Build hourly-stocks-fetcher', style='List Bullet')
    doc.add_paragraph('Build 5min-stocks-top100-fetcher', style='List Bullet')
    doc.add_paragraph('Deploy and test', style='List Bullet')

    add_heading(doc, 'Phase 2: ETF Pipeline (Week 2-3)', 2)
    doc.add_paragraph('Create ETF schema', style='List Bullet')
    doc.add_paragraph('Build ETF fetchers (daily, hourly, 5min)', style='List Bullet')
    doc.add_paragraph('Deploy and test', style='List Bullet')

    add_heading(doc, 'Phase 3: Forex Pipeline (Week 3)', 2)
    doc.add_paragraph('Create forex schema', style='List Bullet')
    doc.add_paragraph('Build forex fetchers (daily, hourly, 5min)', style='List Bullet')
    doc.add_paragraph('Deploy and test', style='List Bullet')

    add_heading(doc, 'Phase 4: Commodities & Bonds (Week 4)', 2)
    doc.add_paragraph('Create commodity/bond schemas', style='List Bullet')
    doc.add_paragraph('Build commodity fetchers', style='List Bullet')
    doc.add_paragraph('Build bond fetchers', style='List Bullet')
    doc.add_paragraph('Deploy and test', style='List Bullet')

    add_heading(doc, 'Phase 5: Fundamental Data (Week 5)', 2)
    doc.add_paragraph('Create fundamental schemas', style='List Bullet')
    doc.add_paragraph('Build fundamental fetchers', style='List Bullet')
    doc.add_paragraph('Deploy weekly jobs', style='List Bullet')
    doc.add_paragraph('Test data quality', style='List Bullet')

    add_heading(doc, 'Phase 6: Frontend Integration (Week 6-7)', 2)
    doc.add_paragraph('Multi-asset selector UI', style='List Bullet')
    doc.add_paragraph('Unified chart component', style='List Bullet')
    doc.add_paragraph('Market overview dashboard', style='List Bullet')
    doc.add_paragraph('Advanced screeners', style='List Bullet')
    doc.add_paragraph('Deploy updated app', style='List Bullet')

    add_heading(doc, 'Phase 7: Testing & Optimization (Week 8)', 2)
    doc.add_paragraph('Load testing', style='List Bullet')
    doc.add_paragraph('Cost optimization', style='List Bullet')
    doc.add_paragraph('Performance tuning', style='List Bullet')
    doc.add_paragraph('Documentation', style='List Bullet')

    doc.add_page_break()

    # Section 10: Asset Universe Details
    add_heading(doc, '10. ASSET UNIVERSE DETAILS', 1)

    add_heading(doc, 'Stocks (S&P 500)', 2)
    doc.add_paragraph('500 largest US companies by market cap')
    doc.add_paragraph('Examples: AAPL, MSFT, GOOGL, AMZN, TSLA, NVDA, META, BRK.B, JNJ, V, WMT, JPM, etc.')

    add_heading(doc, 'ETFs (Top 200)', 2)
    doc.add_paragraph('Sector ETFs: XLF, XLE, XLK, XLV, XLI, XLP, XLU, XLB, XLRE, XLC', style='List Bullet')
    doc.add_paragraph('Index ETFs: SPY, QQQ, IWM, DIA, VOO, VTI, IVV', style='List Bullet')
    doc.add_paragraph('Bond ETFs: AGG, BND, TLT, IEF, SHY, LQD', style='List Bullet')
    doc.add_paragraph('Commodity ETFs: GLD, SLV, USO, UNG', style='List Bullet')
    doc.add_paragraph('International ETFs: EFA, EEM, VWO, IEMG, VEA', style='List Bullet')

    add_heading(doc, 'Forex (Major 20 USD Pairs)', 2)
    doc.add_paragraph('Major: EUR/USD, GBP/USD, JPY/USD, AUD/USD, NZD/USD, CAD/USD, CHF/USD')
    doc.add_paragraph('Emerging: CNY/USD, INR/USD, MXN/USD, BRL/USD, ZAR/USD, TRY/USD, KRW/USD')
    doc.add_paragraph('European: SEK/USD, NOK/USD, DKK/USD, PLN/USD')
    doc.add_paragraph('Asian: HKD/USD, SGD/USD')

    add_heading(doc, 'Crypto (Top 100 USD Pairs)', 2)
    doc.add_paragraph('Large Cap: BTC, ETH, BNB, XRP, ADA, DOGE, SOL, DOT, MATIC')
    doc.add_paragraph('Mid Cap: SHIB, AVAX, UNI, LINK, ATOM, XLM, ALGO, VET, ICP')
    doc.add_paragraph('DeFi: AAVE, SUSHI, COMP, MKR, CRV, YFI, SNX')
    doc.add_paragraph('Layer 1: FIL, EOS, TRX, XTZ, NEO, IOTA')

    add_heading(doc, 'Commodities (40 US-Traded)', 2)
    doc.add_paragraph('Energy: Crude Oil (CL1), Brent (CO1), Natural Gas (NG/USD), Ethanol (DL1)', style='List Bullet')
    doc.add_paragraph('Metals: Gold Gram (GAU/USD), Copper (HG1), Lithium (LC)', style='List Bullet')
    doc.add_paragraph('Agriculture: Corn (C_1), Wheat, Soybeans, Coffee (KC1), Cotton (CT1), Cocoa (CC1)', style='List Bullet')
    doc.add_paragraph('Livestock: Live Cattle (LC1), Feeder Cattle (FC1), Lean Hogs (LH1)', style='List Bullet')
    doc.add_paragraph('Dairy: Milk (DA), Cheese (CHE)', style='List Bullet')
    doc.add_paragraph('Building: Lumber (LB1), Steel (JBP)', style='List Bullet')

    add_heading(doc, 'Bonds (5 US Treasury)', 2)
    doc.add_paragraph('3-Month T-Bill')
    doc.add_paragraph('2-Year Treasury Note')
    doc.add_paragraph('5-Year Treasury Note')
    doc.add_paragraph('10-Year Treasury Note')
    doc.add_paragraph('30-Year Treasury Bond')

    add_heading(doc, 'Indices (5 Major US)', 2)
    doc.add_paragraph('S&P 500 (SPX) - 500 large-cap stocks')
    doc.add_paragraph('Dow Jones Industrial Average (DJI) - 30 blue-chip stocks')
    doc.add_paragraph('Nasdaq Composite (IXIC) - All Nasdaq stocks')
    doc.add_paragraph('Russell 2000 (RUT) - 2000 small-cap stocks')
    doc.add_paragraph('VIX - Volatility Index (Fear Index)')

    doc.add_page_break()

    # Section 11: Security & Best Practices
    add_heading(doc, '11. SECURITY & BEST PRACTICES', 1)

    add_heading(doc, 'API Key Management', 2)
    doc.add_paragraph('Store API key in GCP Secret Manager', style='List Bullet')
    doc.add_paragraph('Rotate API key every 90 days', style='List Bullet')
    doc.add_paragraph('Monitor API usage hourly', style='List Bullet')
    doc.add_paragraph('Set up usage alerts at 80% threshold', style='List Bullet')

    add_heading(doc, 'Rate Limiting', 2)
    doc.add_paragraph('Maximum 1,200 calls/min (75% of limit)', style='List Bullet')
    doc.add_paragraph('Implement exponential backoff on errors', style='List Bullet')
    doc.add_paragraph('Circuit breaker pattern for repeated failures', style='List Bullet')
    doc.add_paragraph('Stagger function execution to avoid conflicts', style='List Bullet')

    add_heading(doc, 'Data Validation', 2)
    doc.add_paragraph('Check for null values before inserting', style='List Bullet')
    doc.add_paragraph('Validate indicator ranges (RSI 0-100, etc.)', style='List Bullet')
    doc.add_paragraph('Duplicate detection using composite keys', style='List Bullet')
    doc.add_paragraph('Data quality reports daily', style='List Bullet')

    add_heading(doc, 'Error Handling', 2)
    doc.add_paragraph('Retry failed symbols up to 3 times', style='List Bullet')
    doc.add_paragraph('Log all errors to BigQuery error table', style='List Bullet')
    doc.add_paragraph('Alert on >10% failure rate', style='List Bullet')
    doc.add_paragraph('Continue processing remaining symbols on failure', style='List Bullet')

    add_heading(doc, 'Monitoring', 2)
    doc.add_paragraph('Cloud Monitoring dashboards for all functions', style='List Bullet')
    doc.add_paragraph('Uptime checks every 5 minutes', style='List Bullet')
    doc.add_paragraph('Cost alerts at $400 monthly threshold', style='List Bullet')
    doc.add_paragraph('Performance metrics tracking', style='List Bullet')

    doc.add_page_break()

    # Section 12: Sample Queries
    add_heading(doc, '12. SAMPLE QUERIES', 1)

    add_heading(doc, 'Find Oversold Stocks with Strong Trends', 2)
    query1 = doc.add_paragraph(
        "SELECT symbol, close, rsi, adx, sma_200\n"
        "FROM `cryptobot-462709.trading_data_unified.stocks_daily`\n"
        "WHERE date = CURRENT_DATE() - 1\n"
        "  AND rsi < 30\n"
        "  AND adx > 25\n"
        "  AND close > sma_200\n"
        "ORDER BY rsi ASC\n"
        "LIMIT 20;"
    )
    query1.style = 'Intense Quote'

    add_heading(doc, 'Top Forex Movers (Last Hour)', 2)
    query2 = doc.add_paragraph(
        "SELECT symbol, close, percent_change\n"
        "FROM `cryptobot-462709.trading_data_unified.forex_hourly`\n"
        "WHERE datetime >= TIMESTAMP_SUB(CURRENT_TIMESTAMP(), INTERVAL 1 HOUR)\n"
        "ORDER BY ABS(percent_change) DESC\n"
        "LIMIT 20;"
    )
    query2.style = 'Intense Quote'

    add_heading(doc, 'Commodities Correlation Matrix', 2)
    query3 = doc.add_paragraph(
        "SELECT\n"
        "  c1.symbol as commodity1,\n"
        "  c2.symbol as commodity2,\n"
        "  CORR(c1.close, c2.close) as correlation\n"
        "FROM `cryptobot-462709.trading_data_unified.commodities_daily` c1\n"
        "JOIN `cryptobot-462709.trading_data_unified.commodities_daily` c2\n"
        "  ON c1.date = c2.date\n"
        "WHERE c1.date >= CURRENT_DATE() - 90\n"
        "  AND c1.symbol < c2.symbol\n"
        "GROUP BY c1.symbol, c2.symbol\n"
        "HAVING correlation > 0.8\n"
        "ORDER BY correlation DESC;"
    )
    query3.style = 'Intense Quote'

    add_heading(doc, 'Stocks Breaking Out of Bollinger Bands', 2)
    query4 = doc.add_paragraph(
        "SELECT symbol, name, close, bb_upper, bb_lower,\n"
        "       (close - bb_upper) / bb_upper * 100 as pct_above_upper\n"
        "FROM `cryptobot-462709.trading_data_unified.stocks_hourly`\n"
        "WHERE datetime >= TIMESTAMP_SUB(CURRENT_TIMESTAMP(), INTERVAL 1 HOUR)\n"
        "  AND close > bb_upper\n"
        "ORDER BY pct_above_upper DESC\n"
        "LIMIT 20;"
    )
    query4.style = 'Intense Quote'

    doc.add_page_break()

    # Section 13: Success Metrics
    add_heading(doc, '13. SUCCESS METRICS', 1)

    add_heading(doc, 'Data Quality Metrics', 2)
    metrics_quality = [
        ['API Success Rate', '99%+', 'Successful API calls'],
        ['Duplicate Rate', '<1%', 'Duplicate records prevented'],
        ['Null Indicator Rate', '<5%', 'Missing indicator values'],
        ['Data Latency', '<5 minutes', 'Time from market to database'],
    ]
    add_table_with_header(doc, ['Metric', 'Target', 'Description'], metrics_quality)

    doc.add_paragraph()
    add_heading(doc, 'Performance Metrics', 2)
    metrics_performance = [
        ['Daily Functions', '<30 min', 'Complete all daily functions'],
        ['Hourly Functions', '<10 min', 'Complete all hourly functions'],
        ['5-Min Functions', '<3 min', 'Complete all 5-min functions'],
        ['Chart Load Time', '<2 seconds', 'Frontend chart rendering'],
    ]
    add_table_with_header(doc, ['Metric', 'Target', 'Description'], metrics_performance)

    doc.add_paragraph()
    add_heading(doc, 'Cost Efficiency Metrics', 2)
    metrics_cost = [
        ['Monthly Cost', '<$500', 'Total infrastructure cost'],
        ['API Utilization', '<15%', 'Percentage of API limit used'],
        ['BigQuery Queries', '<$100/TB', 'Query cost per terabyte'],
        ['Cost per Asset', '<$60/month', 'Cost per asset class'],
    ]
    add_table_with_header(doc, ['Metric', 'Target', 'Description'], metrics_cost)

    doc.add_paragraph()
    add_heading(doc, 'User Engagement Metrics', 2)
    metrics_engagement = [
        ['Frontend Uptime', '99.9%', 'Application availability'],
        ['Real-time Data Delay', '<5 min', 'Data freshness'],
        ['Active Users', 'Track monthly', 'Monthly active users'],
        ['User Retention', '>80%', 'Monthly user retention'],
    ]
    add_table_with_header(doc, ['Metric', 'Target', 'Description'], metrics_engagement)

    doc.add_page_break()

    # Section 14: Maintenance Schedule
    add_heading(doc, '14. MAINTENANCE SCHEDULE', 1)

    add_heading(doc, 'Daily Maintenance', 2)
    doc.add_paragraph('Monitor API usage dashboard', style='List Bullet')
    doc.add_paragraph('Check for failed function executions', style='List Bullet')
    doc.add_paragraph('Review error logs', style='List Bullet')
    doc.add_paragraph('Verify data collection completeness', style='List Bullet')

    add_heading(doc, 'Weekly Maintenance', 2)
    doc.add_paragraph('Analyze cost trends and spending', style='List Bullet')
    doc.add_paragraph('Review data quality metrics', style='List Bullet')
    doc.add_paragraph('Update symbol lists (new IPOs, delistings)', style='List Bullet')
    doc.add_paragraph('Check for API deprecations', style='List Bullet')

    add_heading(doc, 'Monthly Maintenance', 2)
    doc.add_paragraph('Performance optimization review', style='List Bullet')
    doc.add_paragraph('Detailed cost analysis', style='List Bullet')
    doc.add_paragraph('Schema updates if needed', style='List Bullet')
    doc.add_paragraph('User feedback review', style='List Bullet')
    doc.add_paragraph('Backup verification', style='List Bullet')

    add_heading(doc, 'Quarterly Maintenance', 2)
    doc.add_paragraph('API key rotation', style='List Bullet')
    doc.add_paragraph('Major feature updates', style='List Bullet')
    doc.add_paragraph('Infrastructure security audit', style='List Bullet')
    doc.add_paragraph('Disaster recovery testing', style='List Bullet')
    doc.add_paragraph('Performance benchmarking', style='List Bullet')

    doc.add_page_break()

    # Final Section: Support & Next Steps
    add_heading(doc, 'SUPPORT & RESOURCES', 1)

    add_heading(doc, 'Documentation', 2)
    doc.add_paragraph('Twelve Data API Docs: https://twelvedata.com/docs')
    doc.add_paragraph('Twelve Data API Status: https://status.twelvedata.com')
    doc.add_paragraph('GCP Functions Console: https://console.cloud.google.com/functions?project=cryptobot-462709')
    doc.add_paragraph('BigQuery Console: https://console.cloud.google.com/bigquery?project=cryptobot-462709')

    add_heading(doc, 'Support Channels', 2)
    doc.add_paragraph('Twelve Data Support: support@twelvedata.com')
    doc.add_paragraph('GCP Support: Cloud Console Support')
    doc.add_paragraph('Community: Twelve Data GitHub, Stack Overflow')

    doc.add_page_break()

    # Next Steps
    add_heading(doc, 'NEXT STEPS', 1)

    p = doc.add_paragraph()
    p.add_run('1. Review this plan').bold = True
    p.add_run(' - Confirm asset classes and scope\n')

    p.add_run('2. Prioritize phases').bold = True
    p.add_run(' - Which assets to implement first?\n')

    p.add_run('3. Set timeline').bold = True
    p.add_run(' - Target completion date?\n')

    p.add_run('4. Approve budget').bold = True
    p.add_run(' - Is $439/month acceptable?\n')

    p.add_run('5. Begin implementation').bold = True
    p.add_run(' - Start with Phase 1: Stock Data Pipeline\n')

    doc.add_paragraph()
    doc.add_paragraph()

    final = doc.add_paragraph('Ready to build! 🚀')
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final.runs[0].font.size = Pt(18)
    final.runs[0].font.bold = True
    final.runs[0].font.color.rgb = RGBColor(0, 128, 0)

    # Save document
    filename = 'MASTER_TRADING_IMPLEMENTATION_PLAN.docx'
    doc.save(filename)
    print(f"✅ Document created: {filename}")
    return filename

if __name__ == '__main__':
    print("Creating Master Trading Implementation Plan Word document...")
    print("=" * 80)
    filename = create_master_plan_docx()
    print("=" * 80)
    print(f"✅ SUCCESS! Document saved as: {filename}")
    print("📄 Total pages: ~40 pages")
    print("📊 Includes: Tables, formatted headings, sample queries, and complete implementation guide")
