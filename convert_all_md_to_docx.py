"""
Convert ALL Markdown files to DOCX format
"""
import sys
import io
import os
from pathlib import Path
import glob

# Windows encoding fix
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

import re

def parse_markdown_to_docx(md_file, output_file):
    """Convert markdown file to DOCX with formatting"""

    # Read markdown content
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Process line by line
    lines = content.split('\n')
    in_code_block = False
    code_lines = []
    in_list = False
    list_level = 0
    current_table = None

    for line in lines:
        # Skip empty lines in code blocks
        if in_code_block and line.strip() == '':
            code_lines.append('')
            continue

        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'Normal'
                p_format = p.paragraph_format
                p_format.left_indent = Inches(0.5)
                p_format.space_before = Pt(6)
                p_format.space_after = Pt(6)

                # Set background color (light gray)
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)

                code_lines = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
                code_lines = []
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # Skip empty lines
        if not line.strip():
            if not in_list:
                doc.add_paragraph()
            continue

        # Headers
        if line.startswith('#'):
            in_list = False
            current_table = None
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()

            if level == 1:
                p = doc.add_heading(text, level=1)
                if p.runs:
                    p.runs[0].font.size = Pt(24)
                    p.runs[0].font.color.rgb = RGBColor(0, 51, 102)
            elif level == 2:
                p = doc.add_heading(text, level=2)
                if p.runs:
                    p.runs[0].font.size = Pt(18)
                    p.runs[0].font.color.rgb = RGBColor(0, 102, 204)
            elif level == 3:
                p = doc.add_heading(text, level=3)
                if p.runs:
                    p.runs[0].font.size = Pt(14)
                    p.runs[0].font.color.rgb = RGBColor(51, 102, 153)
            else:
                p = doc.add_heading(text, level=4)
                if p.runs:
                    p.runs[0].font.size = Pt(12)
            continue

        # Lists
        if re.match(r'^[\*\-\+]\s+', line) or re.match(r'^\d+\.\s+', line):
            in_list = True
            current_table = None
            # Remove bullet/number
            text = re.sub(r'^[\*\-\+\d\.]+\s+', '', line)

            # Check indentation
            indent = len(line) - len(line.lstrip())
            list_level = indent // 2

            # Add list item
            p = doc.add_paragraph(text, style='List Bullet' if re.match(r'^[\*\-\+]', line.strip()) else 'List Number')
            p.paragraph_format.left_indent = Inches(0.25 * (list_level + 1))
            continue

        # Tables
        if '|' in line and line.strip().startswith('|'):
            in_list = False
            # Skip separator lines
            if re.match(r'^\|[\s\-\|:]+\|$', line.strip()):
                continue

            # Parse table row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]

            # Check if we need to create a new table
            if current_table is None:
                current_table = doc.add_table(rows=1, cols=len(cells))
                current_table.style = 'Light Grid Accent 1'
                is_header = True
            else:
                if is_header:
                    is_header = False
                current_table.add_row()

            # Fill cells
            row = current_table.rows[-1]
            for i, cell_text in enumerate(cells):
                if i < len(row.cells):
                    row.cells[i].text = cell_text
            continue
        else:
            # Reset table tracking when we leave a table
            current_table = None

        # Horizontal rules
        if re.match(r'^[\-\*_]{3,}$', line.strip()):
            in_list = False
            p = doc.add_paragraph()
            p.paragraph_format.border_bottom = True
            continue

        # Regular paragraphs
        in_list = False

        # Process inline formatting
        text = line
        p = doc.add_paragraph()

        # Split by formatting markers
        parts = re.split(r'(\*\*.*?\*\*|`.*?`|\[.*?\]\(.*?\)|🎤|✅|❌|⚠️|🔜|📚|💬|💰|🌍|🌎|🌏|🎯|📊|🔧|📝|🚀|✨|🎉|📄|🔑|💡|⭐|🎨)', text)

        for part in parts:
            if not part:
                continue

            # Bold
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            # Code
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            # Links
            elif '[' in part and '](' in part:
                match = re.match(r'\[(.*?)\]\((.*?)\)', part)
                if match:
                    run = p.add_run(match.group(1))
                    run.font.color.rgb = RGBColor(0, 102, 204)
                    run.underline = True
                else:
                    p.add_run(part)
            # Emojis and regular text
            else:
                p.add_run(part)

    # Save document
    doc.save(output_file)

def main():
    """Main conversion function"""

    print("=" * 80)
    print("MARKDOWN TO DOCX CONVERTER - ALL FILES")
    print("=" * 80)

    # Get script directory
    script_dir = os.path.dirname(os.path.abspath(__file__))

    # Find all .md files
    md_files = glob.glob(os.path.join(script_dir, "*.md"))

    if not md_files:
        print("❌ No markdown files found!")
        return

    print(f"\n📁 Found {len(md_files)} markdown files\n")

    converted = 0
    skipped = 0
    failed = 0

    for md_path in sorted(md_files):
        md_file = os.path.basename(md_path)
        docx_file = md_file.replace('.md', '.docx')
        docx_path = os.path.join(script_dir, docx_file)

        # Check if already exists
        if os.path.exists(docx_path):
            # Check if MD is newer
            md_time = os.path.getmtime(md_path)
            docx_time = os.path.getmtime(docx_path)

            if md_time <= docx_time:
                print(f"⏭️  Skipping {md_file} (DOCX is up-to-date)")
                skipped += 1
                continue

        try:
            print(f"📄 Converting {md_file}...", end=" ")
            parse_markdown_to_docx(md_path, docx_path)
            size_kb = os.path.getsize(docx_path) / 1024
            print(f"✅ ({size_kb:.1f} KB)")
            converted += 1
        except Exception as e:
            print(f"❌ Error: {str(e)}")
            failed += 1

    print("\n" + "=" * 80)
    print("CONVERSION SUMMARY")
    print("=" * 80)
    print(f"✅ Successfully converted: {converted} files")
    if skipped > 0:
        print(f"⏭️  Skipped (up-to-date): {skipped} files")
    if failed > 0:
        print(f"❌ Failed: {failed} files")
    print(f"\n📊 Total: {len(md_files)} markdown files")
    print("=" * 80)

    # List all DOCX files created
    if converted > 0:
        print("\n📂 DOCX files in Trading folder:")
        docx_files = glob.glob(os.path.join(script_dir, "*.docx"))
        for docx_path in sorted(docx_files):
            docx_file = os.path.basename(docx_path)
            size_kb = os.path.getsize(docx_path) / 1024
            print(f"  📄 {docx_file} ({size_kb:.1f} KB)")

if __name__ == "__main__":
    main()
